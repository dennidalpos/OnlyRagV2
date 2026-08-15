import os
import base64
import datetime
import pymupdf
from typing import Dict, Any, List
from sidecar.config import EXPORT_DIR, logger

def _render_pdf_from_markdown(markdown_content: str, output_path: str):
    """Compiles Markdown content to PDF preserving layout, headings, tables, and code blocks."""
    doc = pymupdf.open()
    try:
        page = doc.new_page(width=595, height=842)  # Standard A4 (595x842 pt)
        margin = 40
        page_width = 595 - (2 * margin)
        page_height = 842 - (2 * margin)
        y_offset = margin

        lines = markdown_content.splitlines()
        in_code_block = False
        code_block_lines: List[str] = []

        for raw_line in lines:
            line = raw_line.rstrip()

            # Handle Code Blocks
            if line.strip().startswith("```"):
                if in_code_block:
                    # End of code block - render code box
                    code_text = "\n".join(code_block_lines)
                    code_block_lines = []
                    in_code_block = False
                    
                    # Check page overflow
                    if y_offset + 60 > 842 - margin:
                        page = doc.new_page(width=595, height=842)
                        y_offset = margin

                    rect = pymupdf.Rect(margin, y_offset, margin + page_width, y_offset + 50)
                    page.draw_rect(rect, color=(0.2, 0.2, 0.2), fill=(0.95, 0.95, 0.97))
                    page.insert_textbox(rect, code_text, fontsize=9, fontname="couri", color=(0.1, 0.1, 0.1))
                    y_offset += 60
                else:
                    in_code_block = True
                continue

            if in_code_block:
                code_block_lines.append(line)
                continue

            # Handle Markdown Tables
            if line.strip().startswith("|") and line.strip().endswith("|"):
                parts = [p.strip() for p in line.strip().split("|")[1:-1]]
                if all(set(p) <= {"-", ":", " "} for p in parts if p):
                    continue  # Skip table separator line
                table_row_str = "  |  ".join(parts)
                if y_offset + 18 > 842 - margin:
                    page = doc.new_page(width=595, height=842)
                    y_offset = margin
                rect = pymupdf.Rect(margin, y_offset, margin + page_width, y_offset + 16)
                page.draw_rect(rect, color=(0.8, 0.8, 0.8), fill=(0.96, 0.96, 0.96))
                page.insert_textbox(rect, table_row_str, fontsize=10, fontname="helv", color=(0.1, 0.1, 0.3))
                y_offset += 18
                continue

            if not line.strip():
                y_offset += 10
                if y_offset > 842 - margin:
                    page = doc.new_page(width=595, height=842)
                    y_offset = margin
                continue

            # Determine typography style
            if line.startswith("# "):
                font_size = 18
                line_height = 24
                font_name = "hebo"
                clean_text = line[2:].strip()
                color = (0.05, 0.15, 0.35)
            elif line.startswith("## "):
                font_size = 14
                line_height = 20
                font_name = "hebo"
                clean_text = line[3:].strip()
                color = (0.1, 0.25, 0.45)
            elif line.startswith("### "):
                font_size = 12
                line_height = 16
                font_name = "hebo"
                clean_text = line[4:].strip()
                color = (0.15, 0.3, 0.5)
            elif line.startswith("* ") or line.startswith("- "):
                font_size = 10
                line_height = 14
                font_name = "helv"
                clean_text = "• " + line[2:].strip()
                color = (0.1, 0.1, 0.1)
            else:
                font_size = 10
                line_height = 14
                font_name = "helv"
                clean_text = line.strip()
                color = (0.1, 0.1, 0.1)

            # Check page boundary
            if y_offset + line_height > 842 - margin:
                page = doc.new_page(width=595, height=842)
                y_offset = margin

            rect = pymupdf.Rect(margin, y_offset, margin + page_width, y_offset + line_height + 4)
            page.insert_textbox(rect, clean_text, fontsize=font_size, fontname=font_name, color=color)
            y_offset += line_height

        doc.save(output_path)
    finally:
        doc.close()


def _render_docx_from_markdown(markdown_content: str, output_path: str):
    """Compiles Markdown content to DOCX Word document maintaining headings, lists, tables, and formatting."""
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = docx.Document()
    
    # Configure page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = docx.shared.Inches(0.75)
        section.bottom_margin = docx.shared.Inches(0.75)
        section.left_margin = docx.shared.Inches(0.75)
        section.right_margin = docx.shared.Inches(0.75)

    lines = markdown_content.splitlines()
    in_code_block = False
    code_block_lines: List[str] = []
    table_rows_buffer: List[List[str]] = []

    def flush_table_buffer():
        nonlocal table_rows_buffer
        if not table_rows_buffer:
            return
        
        cols_count = max(len(r) for r in table_rows_buffer)
        table = doc.add_table(rows=len(table_rows_buffer), cols=cols_count)
        table.style = 'Table Grid'
        
        for r_idx, row in enumerate(table_rows_buffer):
            for c_idx, cell_value in enumerate(row):
                if c_idx < cols_count:
                    cell = table.cell(r_idx, c_idx)
                    cell.text = cell_value
                    if r_idx == 0:
                        # Bold table headers
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.bold = True
        table_rows_buffer = []

    for raw_line in lines:
        line = raw_line.rstrip()

        # Handle Code Blocks
        if line.strip().startswith("```"):
            if in_code_block:
                flush_table_buffer()
                code_text = "\n".join(code_block_lines)
                code_block_lines = []
                in_code_block = False
                
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(12)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(40, 40, 40)
            else:
                flush_table_buffer()
                in_code_block = True
            continue

        if in_code_block:
            code_block_lines.append(line)
            continue

        # Handle Markdown Tables
        if line.strip().startswith("|") and line.strip().endswith("|"):
            parts = [p.strip() for p in line.strip().split("|")[1:-1]]
            if all(set(p) <= {"-", ":", " "} for p in parts if p):
                continue
            table_rows_buffer.append(parts)
            continue
        else:
            flush_table_buffer()

        if not line.strip():
            continue

        # Headings and Text Formatting
        if line.startswith("# "):
            h = doc.add_heading(line[2:].strip(), level=1)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
        elif line.startswith("## "):
            h = doc.add_heading(line[3:].strip(), level=2)
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(4)
        elif line.startswith("### "):
            h = doc.add_heading(line[4:].strip(), level=3)
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(2)
        elif line.startswith("* ") or line.startswith("- "):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(line[2:].strip())
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            p.add_run(line.strip())

    flush_table_buffer()
    doc.save(output_path)


def export_markdown_to_file(markdown_content: str, export_format: str = "pdf") -> Dict[str, Any]:
    """Compiles Markdown into structured PDF, DOCX, HTML, or raw MD file preserving original layout."""
    fmt = export_format.lower().strip()
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

    if fmt == "pdf":
        export_filename = f"export_{ts}.pdf"
        file_path = os.path.join(EXPORT_DIR, export_filename)
        try:
            _render_pdf_from_markdown(markdown_content, file_path)
            with open(file_path, "rb") as f:
                b64_content = base64.b64encode(f.read()).decode("utf-8")

            return {
                "status": "success",
                "format": "pdf",
                "file_name": export_filename,
                "file_path": file_path,
                "base64_content": b64_content,
                "message": f"Successfully compiled markdown into PDF format: {export_filename}"
            }
        except Exception as e:
            logger.error(f"Failed PDF compilation in sidecar exporter: {e}")
            raise RuntimeError(f"PDF export failed: {str(e)}")

    elif fmt == "docx":
        export_filename = f"export_{ts}.docx"
        file_path = os.path.join(EXPORT_DIR, export_filename)
        try:
            _render_docx_from_markdown(markdown_content, file_path)
            with open(file_path, "rb") as f:
                b64_content = base64.b64encode(f.read()).decode("utf-8")

            return {
                "status": "success",
                "format": "docx",
                "file_name": export_filename,
                "file_path": file_path,
                "base64_content": b64_content,
                "message": f"Successfully compiled markdown into DOCX format: {export_filename}"
            }
        except Exception as e:
            logger.error(f"Failed DOCX compilation in sidecar exporter: {e}")
            raise RuntimeError(f"DOCX export failed: {str(e)}")

    elif fmt in ["html", "htm"]:
        export_filename = f"export_{ts}.html"
        file_path = os.path.join(EXPORT_DIR, export_filename)
        try:
            html_wrapper = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Exported Document</title>
    <style>
        body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; line-height: 1.6; max-width: 850px; margin: 40px auto; padding: 0 20px; color: #1e293b; background: #0f172a; color: #f8fafc; }}
        h1, h2, h3 {{ color: #38bdf8; }}
        pre {{ background: #1e293b; padding: 16px; border-radius: 8px; overflow-x: auto; font-family: monospace; }}
        table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
        th, td {{ border: 1px solid #334155; padding: 8px 12px; text-align: left; }}
        th {{ background: #1e293b; color: #38bdf8; }}
    </style>
</head>
<body>
    <pre>{markdown_content}</pre>
</body>
</html>"""
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(html_wrapper)

            with open(file_path, "rb") as f:
                b64_content = base64.b64encode(f.read()).decode("utf-8")

            return {
                "status": "success",
                "format": "html",
                "file_name": export_filename,
                "file_path": file_path,
                "base64_content": b64_content,
                "message": f"Successfully exported HTML document: {export_filename}"
            }
        except Exception as e:
            logger.error(f"Failed HTML export in sidecar exporter: {e}")
            raise RuntimeError(f"HTML export failed: {str(e)}")

    else:
        export_filename = f"export_{ts}.md"
        file_path = os.path.join(EXPORT_DIR, export_filename)
        try:
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(markdown_content)

            with open(file_path, "rb") as f:
                b64_content = base64.b64encode(f.read()).decode("utf-8")

            return {
                "status": "success",
                "format": fmt,
                "file_name": export_filename,
                "file_path": file_path,
                "base64_content": b64_content,
                "message": f"Successfully exported markdown file: {export_filename}"
            }
        except Exception as e:
            logger.error(f"Failed file export in sidecar exporter: {e}")
            raise RuntimeError(f"Export failed: {str(e)}")
