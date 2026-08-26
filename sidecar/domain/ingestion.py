import os
import re
import io
import json
import csv
import zipfile
from typing import List, Tuple, Dict, Optional, Any, Callable
from concurrent.futures import ThreadPoolExecutor
import pymupdf
import pandas as pd
from sidecar.config import (
    logger,
    OLLAMA_BASE_URL,
    TABULAR_MAX_ROWS,
    EXCEL_MAX_ROWS_PER_SHEET,
    EXCEL_MAX_SHEETS
)
from sidecar.infrastructure.ocr import run_layout_ocr, run_vision_ocr
from sidecar.domain.sanitizer import sanitize_extracted_text
from sidecar.domain.vision_prompt import is_vision_ocr_requested, render_vision_prompt
from sidecar.domain.router import (
    classify_file_type,
    DocumentCategory,
    analyze_pdf_page_structure,
    PageRoutingStrategy
)

# Bounded concurrency for the OCR/Vision rendering phase of PDF page extraction. Kept small and
# shared for both the local RapidOCR tier and the Ollama Vision fallback: RapidOCR could safely
# take more, but a batch may include several pages that fall back to the single shared local
# Ollama daemon, and overloading it risks the VRAM thrashing the rest of the architecture avoids.
PDF_PAGE_RENDER_CONCURRENCY = 3

_FIGURE_VISION_PROMPT = (
    "Describe this figure, chart, table, or diagram concisely in Markdown: "
    "extract any visible text, labels, axis values, and data points. "
    "If it is a purely decorative image with no informational content, respond with an empty string."
)

def extract_tables_from_page(page: pymupdf.Page) -> Tuple[List[str], List[Any]]:
    """Extracts tables natively from a PDF page and converts them to formatted Markdown tables via tabulate."""
    md_tables: List[str] = []
    table_rects: List[Any] = []
    try:
        tabs = page.find_tables()
        if tabs and tabs.tables:
            for tab in tabs.tables:
                table_rects.append(tab.bbox)
                try:
                    df = tab.to_pandas()
                    if df is not None and not df.empty:
                        df.columns = [str(c).replace("\n", " ").strip() if str(c).strip() else f"Col {i+1}" for i, c in enumerate(df.columns)]
                        md_table = df.to_markdown(tablefmt="pipe", index=False)
                        if md_table:
                            md_tables.append("\n" + md_table + "\n")
                            continue
                except Exception:
                    pass

                rows = tab.extract()
                if not rows or len(rows) < 1:
                    continue

                clean_rows = []
                for r in rows:
                    if r and any(cell and str(cell).strip() for cell in r):
                        clean_rows.append([str(c or '').replace('|', '\\|').replace('\n', ' ').strip() for c in r])

                if not clean_rows:
                    continue

                headers = clean_rows[0]
                safe_headers = [h if h else f"Col {i+1}" for i, h in enumerate(headers)]
                df_fallback = pd.DataFrame(clean_rows[1:], columns=safe_headers)
                md_fallback = df_fallback.to_markdown(tablefmt="pipe", index=False)
                if md_fallback:
                    md_tables.append("\n" + md_fallback + "\n")
    except Exception as e:
        logger.debug(f"Table detection skipped or not supported on page: {e}")

    return md_tables, table_rects

from sidecar.domain.llm_normalizer import normalize_page_markdown_with_llm

def run_page_ocr(
    image_bytes: bytes,
    vision_model: Optional[str] = None,
    vision_prompt: Optional[str] = None,
    filename: str = "",
    page_num: int = 1,
    num_pages: int = 1,
    active_page_content: str = "",
    num_ctx: Optional[int] = None
) -> str:
    """
    Routes one page/image bitmap to the OCR engine the user selected.

    A non-empty `vision_prompt` means the multimodal Vision LLM engine was chosen
    (`AppSettings.ocrEngine === 'vision_model'`); RapidOCR then stays as the safety net, because a
    vision call returns an empty string on a missing model, a timeout or a refusal, and a page must
    never be silently dropped over that.
    """
    if is_vision_ocr_requested(vision_prompt):
        prompt = render_vision_prompt(vision_prompt, filename, page_num, num_pages, active_page_content)
        vision_kwargs = {
            "ollama_url": OLLAMA_BASE_URL,
            "model": vision_model or "llama3.2-vision",
        }
        if num_ctx is not None:
            vision_kwargs["num_ctx"] = num_ctx
        vision_text = run_vision_ocr(image_bytes, prompt, **vision_kwargs)
        if vision_text.strip():
            return vision_text
        logger.warning(f"Vision OCR yielded no content on page {page_num}; falling back to RapidOCR.")

    return run_layout_ocr(image_bytes)

def prepare_pdf_page_work_item(
    doc: pymupdf.Document,
    page: pymupdf.Page,
    page_num: int,
    raw_text: str,
    md_tables: List[str],
    used_ocr: bool,
    vision_model: Optional[str] = None,
    vision_prompt: Optional[str] = None,
    normalize_with_llm: bool = False,
    normalization_model: Optional[str] = None,
    filename: str = "",
    num_pages: int = 1,
    num_ctx: Optional[int] = None,
    **kwargs: Any
) -> Dict[str, Any]:
    """
    Sequential, PyMuPDF-bound preparation step for a single PDF page: renders the page image for
    OCR (if used_ocr or native text is sparse). Must run on the thread that owns doc/page.
    """
    ocr_image_bytes: Optional[bytes] = None
    if used_ocr or len(raw_text.strip()) == 0:
        try:
            pix = page.get_pixmap(dpi=250)
            ocr_image_bytes = pix.tobytes("png")
        except Exception as pix_err:
            logger.debug(f"Pixmap generation skipped on page {page_num}: {pix_err}")

    return {
        "page_num": page_num,
        "raw_text": raw_text,
        "md_tables": md_tables,
        "used_ocr": used_ocr,
        "ocr_image_bytes": ocr_image_bytes,
        "vision_model": vision_model,
        "vision_prompt": vision_prompt,
        "normalize_with_llm": normalize_with_llm,
        "normalization_model": normalization_model,
        "filename": filename,
        "num_pages": num_pages,
        "num_ctx": num_ctx,
    }

def _ocr_work_item_image(work_item: Dict[str, Any], active_page_content: str = "") -> str:
    """Unpacks a prepared work item onto `run_page_ocr`."""
    return run_page_ocr(
        work_item["ocr_image_bytes"],
        vision_model=work_item.get("vision_model"),
        vision_prompt=work_item.get("vision_prompt"),
        filename=work_item.get("filename", ""),
        page_num=work_item["page_num"],
        num_pages=work_item.get("num_pages", 1),
        active_page_content=active_page_content,
        num_ctx=work_item.get("num_ctx")
    )

def render_prepared_pdf_page(work_item: Dict[str, Any]) -> Tuple[int, str]:
    """
    Renders a single prepared page: runs the selected OCR engine (see `run_page_ocr`) on
    scanned/bitmap pages, or formats native text/tables. Applies optional LLM Markdown
    normalization only if explicitly requested.
    """
    page_num = work_item["page_num"]
    page_md_parts: List[str] = []

    if work_item["used_ocr"] and work_item.get("ocr_image_bytes"):
        ocr_result = _ocr_work_item_image(work_item, active_page_content=work_item.get("raw_text", ""))
        if ocr_result.strip():
            page_md_parts.append(ocr_result.strip())
        else:
            page_md_parts.append("[Scanned page - No readable text detected]")
    else:
        native_text = (work_item.get("raw_text") or "").strip()
        if native_text:
            page_md_parts.append(native_text)
        if work_item.get("md_tables"):
            page_md_parts.extend(work_item["md_tables"])

        # Fallback for scanned pages where native text layer was empty
        if not native_text and not work_item.get("md_tables") and work_item.get("ocr_image_bytes"):
            ocr_result = _ocr_work_item_image(work_item)
            if ocr_result.strip():
                page_md_parts.append(ocr_result.strip())

    page_content = "\n\n".join(page_md_parts).strip()
    if not page_content:
        page_content = "[Empty Page Content]"

    sanitized = sanitize_extracted_text(page_content)

    if work_item.get("normalize_with_llm"):
        norm_model = work_item.get("normalization_model") or "llama3.2"
        sanitized = normalize_page_markdown_with_llm(sanitized, page_num=page_num, model=norm_model)

    return page_num, sanitized

def extract_pdf_document(
    doc: pymupdf.Document,
    progress_callback: Optional[Callable[[int, int, str], None]] = None,
    vision_model: Optional[str] = None,
    vision_prompt: Optional[str] = None,
    normalize_with_llm: bool = False,
    normalization_model: Optional[str] = None,
    filename: str = "",
    num_ctx: Optional[int] = None
) -> List[Tuple[int, str]]:
    """
    Extracts markdown per page preserving layout, native tables, and high-fidelity OCR.
    PyMuPDF access happens sequentially; OCR rendering runs concurrently in a bounded thread pool.
    """
    num_pages = len(doc)
    work_items: List[Dict[str, Any]] = []

    for page_idx in range(num_pages):
        page_num = page_idx + 1
        page = doc.load_page(page_idx)

        struct_info = analyze_pdf_page_structure(page)
        strategy = struct_info.get("strategy")

        md_tables, _ = extract_tables_from_page(page)
        table_info = f" (trovate {len(md_tables)} tabelle)" if md_tables else ""

        if progress_callback:
            progress_callback(
                page_num,
                num_pages,
                f"Elaborazione pagina {page_num}/{num_pages}{table_info}..."
            )

        raw_text = page.get_text("text").strip()
        used_ocr = strategy == PageRoutingStrategy.OCR_REQUIRED

        if used_ocr and progress_callback:
            progress_callback(page_num, num_pages, f"Pagina {page_num}/{num_pages}: Esecuzione OCR Layout...")

        work_items.append(prepare_pdf_page_work_item(
            doc, page, page_num, raw_text, md_tables, used_ocr,
            vision_model=vision_model,
            vision_prompt=vision_prompt,
            normalize_with_llm=normalize_with_llm,
            normalization_model=normalization_model,
            filename=filename,
            num_pages=num_pages,
            num_ctx=num_ctx
        ))

    with ThreadPoolExecutor(max_workers=min(PDF_PAGE_RENDER_CONCURRENCY, max(1, len(work_items)))) as executor:
        results = list(executor.map(render_prepared_pdf_page, work_items))

    return results

def extract_tabular_document(
    filename: str,
    content: bytes,
    file_path: Optional[str],
    max_rows: Optional[int] = None,
    max_excel_rows: Optional[int] = None,
    max_sheets: Optional[int] = None,
    num_ctx: Optional[int] = None
) -> List[Tuple[int, str]]:
    """Extracts CSV, TSV, XLSX, XLS, Parquet, and JSON data formatted cleanly into Markdown tables with truncation tracking."""
    ext = os.path.splitext(filename)[1].lower()
    text_content = ""
    limit_csv_rows = max_rows or TABULAR_MAX_ROWS
    limit_excel_rows = max_excel_rows or EXCEL_MAX_ROWS_PER_SHEET
    limit_sheets = max_sheets or EXCEL_MAX_SHEETS

    # Excel formats (.xlsx, .xls)
    if ext in [".xlsx", ".xls"]:
        try:
            excel_source = file_path if (file_path and os.path.exists(file_path)) else io.BytesIO(content)
            xls = pd.ExcelFile(excel_source)
            sheets_output = []
            total_sheets_count = len(xls.sheet_names)
            processed_sheets = xls.sheet_names[:limit_sheets]

            for sheet_name in processed_sheets:
                df_all = pd.read_excel(xls, sheet_name=sheet_name)
                total_sheet_rows = len(df_all)
                if df_all.empty:
                    continue
                df = df_all.head(limit_excel_rows)
                # Clean column headers
                df.columns = [str(c).replace("\n", " ").strip() if str(c).strip() else f"Col {i+1}" for i, c in enumerate(df.columns)]
                md_table = df.to_markdown(tablefmt="pipe", index=False) or ""

                sheet_md = f"### Foglio: {sheet_name}\n\n{md_table}"
                if total_sheet_rows > limit_excel_rows:
                    sheet_md += f"\n\n> [!NOTE]\n> Tabella troncata a {limit_excel_rows} righe (su {total_sheet_rows} totali nel foglio '{sheet_name}').\n"
                sheets_output.append(sheet_md)

            combined_excel_md = "\n\n".join(sheets_output)
            if total_sheets_count > limit_sheets:
                combined_excel_md += f"\n\n> [!NOTE]\n> Mostrati i primi {len(processed_sheets)} fogli su {total_sheets_count} totali nella cartella Excel.\n"

            return [(1, combined_excel_md)]
        except Exception as e:
            logger.warning(f"Excel parsing fallback for {filename}: {e}")

    # CSV / TSV
    if ext in [".csv", ".tsv"]:
        try:
            csv_source = file_path if (file_path and os.path.exists(file_path)) else io.BytesIO(content)
            sep = "\t" if ext == ".tsv" else ","
            df_all = pd.read_csv(csv_source, sep=sep, on_bad_lines="skip")
            total_csv_rows = len(df_all)
            df = df_all.head(limit_csv_rows)
            df.columns = [str(c).replace("\n", " ").strip() if str(c).strip() else f"Col {i+1}" for i, c in enumerate(df.columns)]
            md_table = df.to_markdown(tablefmt="pipe", index=False) or ""

            if total_csv_rows > limit_csv_rows:
                md_table += f"\n\n> [!NOTE]\n> Tabella troncata a {limit_csv_rows} righe (su {total_csv_rows} totali nel file).\n"

            return [(1, md_table)]
        except Exception as e:
            logger.warning(f"CSV/TSV parsing fallback for {filename}: {e}")

    # JSON formatted
    if ext == ".json":
        try:
            if file_path and os.path.exists(file_path):
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    text_content = f.read()
            else:
                text_content = content.decode("utf-8", errors="ignore")
            data = json.loads(text_content)
            pretty_json = json.dumps(data, indent=2, ensure_ascii=False)
            return [(1, f"```json\n{pretty_json}\n```")]
        except Exception:
            pass

    return [(1, text_content.strip())]

def extract_document_markdown(
    filename: str,
    content: bytes,
    file_path: Optional[str] = None,
    progress_callback: Optional[Callable[[int, int, str], None]] = None,
    vision_model: Optional[str] = None,
    vision_prompt: Optional[str] = None,
    normalize_with_llm: bool = False,
    normalization_model: Optional[str] = None,
    max_tabular_rows: Optional[int] = None,
    max_excel_rows: Optional[int] = None,
    max_sheets: Optional[int] = None,
    num_ctx: Optional[int] = None
) -> Tuple[str, int]:
    """Fast-routed, sanitized, and pagination-preserving document markdown extractor with progress callback."""
    category = classify_file_type(filename)
    num_pages = 1
    page_blocks: List[Tuple[int, str]] = []

    if category == DocumentCategory.PDF:
        if file_path and os.path.exists(file_path):
            pdf_doc = pymupdf.open(file_path)
        else:
            pdf_doc = pymupdf.open(stream=content, filetype="pdf")

        if pdf_doc.needs_pass != 0 or (pdf_doc.is_encrypted and pdf_doc.needs_pass):
            pdf_doc.close()
            raise ValueError("Documento protetto da password: il file PDF è crittografato e richiede una password per l'apertura.")

        try:
            num_pages = len(pdf_doc)
            page_blocks = extract_pdf_document(
                pdf_doc,
                progress_callback=progress_callback,
                vision_model=vision_model,
                vision_prompt=vision_prompt,
                normalize_with_llm=normalize_with_llm,
                normalization_model=normalization_model,
                filename=filename,
                num_ctx=num_ctx
            )
        finally:
            pdf_doc.close()

    elif category == DocumentCategory.IMAGE:
        if progress_callback:
            progress_callback(1, 1, "Esecuzione OCR su immagine...")
        img_bytes = content
        if file_path and os.path.exists(file_path):
            with open(file_path, "rb") as f:
                img_bytes = f.read()
        ocr_text = run_page_ocr(
            img_bytes,
            vision_model=vision_model,
            vision_prompt=vision_prompt,
            filename=filename,
            num_ctx=num_ctx
        )
        sanitized_ocr = sanitize_extracted_text(ocr_text)
        if normalize_with_llm and sanitized_ocr:
            norm_model = normalization_model or "llama3.2"
            sanitized_ocr = normalize_page_markdown_with_llm(sanitized_ocr, page_num=1, model=norm_model)
        if sanitized_ocr:
            page_blocks.append((1, sanitized_ocr))
        else:
            page_blocks.append((1, f"[Image file {filename} scanned - No text detected]"))

    elif category == DocumentCategory.DOCX:
        if progress_callback:
            progress_callback(1, 1, "Estrazione struttura XML, tabelle e immagini DOCX...")
        try:
            import docx
            # Check for EncryptedPackage OLE header
            if file_path and os.path.exists(file_path):
                with open(file_path, "rb") as f:
                    hdr = f.read(2048)
            else:
                hdr = content[:2048]

            if hdr.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1') and (b'EncryptedPackage' in hdr or b'EncryptionInfo' in hdr):
                raise ValueError("Documento protetto da password: il file Word (DOCX) è crittografato e richiede una password.")

            if file_path and os.path.exists(file_path):
                doc = docx.Document(file_path)
            else:
                doc = docx.Document(io.BytesIO(content))

            docx_blocks: List[str] = []
            for para in doc.paragraphs:
                p_text = para.text.strip()
                if not p_text:
                    continue
                style_name = (para.style.name if para.style else "").lower()
                if "heading 1" in style_name:
                    docx_blocks.append(f"# {p_text}")
                elif "heading 2" in style_name:
                    docx_blocks.append(f"## {p_text}")
                elif "heading 3" in style_name:
                    docx_blocks.append(f"### {p_text}")
                else:
                    docx_blocks.append(p_text)

            for table in doc.tables:
                table_md_lines: List[str] = []
                headers_done = False
                for row in table.rows:
                    row_cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                    if not any(row_cells):
                        continue
                    table_md_lines.append("| " + " | ".join(row_cells) + " |")
                    if not headers_done:
                        table_md_lines.append("| " + " | ".join(["---"] * len(row_cells)) + " |")
                        headers_done = True
                if table_md_lines:
                    docx_blocks.append("\n" + "\n".join(table_md_lines) + "\n")

            # Extract embedded images from docx package and run RapidOCR
            try:
                docx_zip_source = file_path if (file_path and os.path.exists(file_path)) else io.BytesIO(content)
                with zipfile.ZipFile(docx_zip_source, "r") as z:
                    image_names = sorted([n for n in z.namelist() if n.startswith("word/media/") and any(n.lower().endswith(im_ext) for im_ext in [".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"])])
                    img_idx = 1
                    for img_name in image_names:
                        img_data = z.read(img_name)
                        if len(img_data) >= 1500:  # Skip tiny icons or empty placeholder graphics
                            ocr_text = run_page_ocr(
                                img_data,
                                vision_model=vision_model,
                                vision_prompt=vision_prompt,
                                filename=filename,
                                num_ctx=num_ctx
                            )
                            clean_ocr = sanitize_extracted_text(ocr_text)
                            if clean_ocr and clean_ocr not in ("[Empty Page Content]", "[Scanned page - No readable text detected]"):
                                docx_blocks.append(f"### Immagine DOCX {img_idx} (Testo rilevato da OCR)\n\n{clean_ocr}")
                                img_idx += 1
            except Exception as zip_err:
                logger.debug(f"DOCX media image extraction skipped for {filename}: {zip_err}")

            if docx_blocks:
                page_blocks.append((1, sanitize_extracted_text("\n\n".join(docx_blocks))))
        except Exception as docx_err:
            logger.warning(f"python-docx parse error for {filename}: {docx_err}")

    elif category == DocumentCategory.TABULAR:
        if progress_callback:
            progress_callback(1, 1, "Parsing matriciale e formattazione tabella Markdown...")
        page_blocks = extract_tabular_document(
            filename,
            content,
            file_path,
            max_rows=max_tabular_rows,
            max_excel_rows=max_excel_rows,
            max_sheets=max_sheets
        )

    # Fallback to UTF-8 text read
    if not page_blocks:
        if progress_callback:
            progress_callback(1, 1, "Lettura e sanitizzazione flusso testuale...")
        try:
            if file_path and os.path.exists(file_path):
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    raw_str = f.read()
            else:
                raw_str = content.decode("utf-8", errors="ignore")
            page_blocks.append((1, sanitize_extracted_text(raw_str)))
        except Exception:
            page_blocks.append((1, "[Binary content processed]"))

    # Assemble strictly paginated Markdown document
    paginated_sections: List[str] = []
    for page_idx, p_text in page_blocks:
        paginated_sections.append(f"## Page {page_idx}\n\n{p_text}")

    full_markdown = f"# {filename}\n\n" + "\n\n".join(paginated_sections)
    return full_markdown, max(1, num_pages)

def _split_oversized_text(text: str, max_chars: int = 1000, overlap: int = 100) -> List[str]:
    """Splits a single oversized paragraph or text block at whitespace or sentence boundaries."""
    if len(text) <= max_chars:
        return [text]
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + max_chars, len(text))
        if end < len(text):
            # Look for sentence or whitespace boundary
            boundary = max(text.rfind('. ', start, end), text.rfind('\n', start, end), text.rfind(' ', start, end))
            if boundary > start + (max_chars // 2):
                end = boundary + 1
        chunks.append(text[start:end].strip())
        start = end - overlap if end < len(text) else end
    return [c for c in chunks if c]

def create_semantic_chunks(filename: str, full_markdown: str) -> List[Tuple[int, str, str]]:
    """
    Header-Aware Semantic Chunking:
    Preserves Markdown structural hierarchy, section context prefixes, and code blocks.
    Returns list of (chunk_index, chunk_text, section_header).
    """
    raw_chunks: List[Tuple[int, str, str]] = []
    header_path: List[str] = [filename]
    chunk_index = 0
    current_buffer: List[str] = []
    in_code_block = False

    def _flush_buffer():
        nonlocal chunk_index, current_buffer
        if not current_buffer:
            return
        curr_section_header = ' > '.join([h for h in header_path if h])
        context_prefix = f"[Documento: {filename} | Sezione: {curr_section_header}]\n"
        buf_text = "\n".join(current_buffer).strip()
        current_buffer = []
        if not buf_text:
            return

        if len(buf_text) > 800:
            pieces = _split_oversized_text(buf_text, max_chars=800, overlap=100)
            for p in pieces:
                if p.strip():
                    raw_chunks.append((chunk_index, context_prefix + p.strip(), curr_section_header))
                    chunk_index += 1
        else:
            raw_chunks.append((chunk_index, context_prefix + buf_text, curr_section_header))
            chunk_index += 1

    lines = full_markdown.splitlines()
    for line in lines:
        stripped_line = line.strip()

        if stripped_line.startswith("```"):
            in_code_block = not in_code_block

        if not in_code_block and stripped_line.startswith("#"):
            header_match = re.match(r'^(#+)\s+(.+)$', stripped_line)
            if header_match:
                _flush_buffer()
                level = len(header_match.group(1))
                h_text = header_match.group(2).strip()
                header_path = header_path[:level]
                if len(header_path) < level:
                    header_path.extend([""] * (level - len(header_path)))
                header_path = header_path[:level-1] + [h_text]
                continue

        current_buffer.append(line)

    _flush_buffer()

    if not raw_chunks:
        fallback_text = full_markdown.strip() or f"# {filename}\n\nDocument content ingested."
        raw_chunks = [(0, f"[Documento: {filename} | Sezione: Document Content]\n{fallback_text}", filename)]

    return raw_chunks

