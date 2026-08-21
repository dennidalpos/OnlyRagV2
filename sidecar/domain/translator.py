import os
import re
import time
import uuid
from typing import Any, Dict, List, Tuple, Optional
import docx
import httpx
import pymupdf
from sidecar.config import DOCS_TABLE_NAME, logger, httpx_client
from sidecar.infrastructure.db import lance_db, get_existing_tables, validate_doc_id
from sidecar.schemas import IngestResponse
from sidecar.domain.ingestion import extract_document_markdown
from sidecar.services.ingest_service import update_and_reindex_document

# Delimiter injected between run texts in a translation batch prompt. The model is instructed to
# preserve it verbatim so the response can be split back into the same number of segments, in the
# same order, and reassigned 1:1 to the runs that produced them.
_RUN_SEPARATOR = "<<<RUN_SEP>>>"
_TRANSLATE_BATCH_MAX_CHARS = 350
_TRANSLATE_BATCH_MAX_ITEMS = 4
_OLLAMA_URL = "http://127.0.0.1:11434"

# Fase 3 auto-fit: never shrink text below this absolute size or this fraction of the original
# span size, whichever floor is higher -- keeps reinserted text legible instead of vanishing.
_PDF_AUTOFIT_MIN_SIZE = 6.0
_PDF_AUTOFIT_MIN_RATIO = 0.4
# Binary search stops refining once the [lo, hi] font-size bracket is this narrow.
_PDF_AUTOFIT_TOLERANCE = 0.25

# Fase 4: bundled static Regular-weight fonts (see sidecar/assets/fonts/OFL-*.txt for license and
# provenance -- SIL Open Font License 1.1, Noto Project). No PDF base14 font covers CJK or
# Cyrillic/Greek, so every PDF reinsertion goes through one of these instead of a built-in font.
_PDF_FONT_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "assets", "fonts")
_PDF_FALLBACK_FONT_FILE = os.path.join(_PDF_FONT_DIR, "NotoSans-Regular.otf")  # Latin, Cyrillic, Greek

# target_lang (free-text from the frontend language picker) -> bundled font file, matched
# case-insensitively by substring, first match wins. More specific hints are listed before the
# broader ones they'd otherwise be shadowed by (e.g. "traditional chinese" before "chinese").
# Anything not matched here falls back to _PDF_FALLBACK_FONT_FILE.
_PDF_LANG_FONT_RULES: List[Tuple[str, str]] = [
    ("japanese", os.path.join(_PDF_FONT_DIR, "NotoSansCJKjp-Regular.otf")),
    ("korean", os.path.join(_PDF_FONT_DIR, "NotoSansCJKkr-Regular.otf")),
    ("traditional chinese", os.path.join(_PDF_FONT_DIR, "NotoSansCJKtc-Regular.otf")),
    ("chinese", os.path.join(_PDF_FONT_DIR, "NotoSansCJKsc-Regular.otf")),  # default: Simplified
]


_LANG_PATTERNS = {
    "chinese": re.compile(r'[\u4e00-\u9fff]'),
    "japanese": re.compile(r'[\u3040-\u30ff\u31f0-\u31ff]'),
    "korean": re.compile(r'[\uac00-\ud7af\u1100-\u11ff]'),
    "cyrillic": re.compile(r'[\u0400-\u04ff]'),
    "arabic": re.compile(r'[\u0600-\u06ff]'),
    "greek": re.compile(r'[\u0370-\u03ff]'),
}

_LANGDETECT_AVAILABLE = False
try:
    import langdetect
    from langdetect import DetectorFactory
    DetectorFactory.seed = 0
    _LANGDETECT_AVAILABLE = True
except ImportError:
    _LANGDETECT_AVAILABLE = False

_ISO_TO_LANG_NAME: Dict[str, str] = {
    "it": "italian", "en": "english", "es": "spanish", "fr": "french",
    "de": "german", "pt": "portuguese", "nl": "dutch", "ru": "russian",
    "zh-cn": "chinese", "zh-tw": "chinese", "zh": "chinese", "ja": "japanese",
    "ko": "korean", "ar": "arabic", "el": "greek", "pl": "polish",
    "sv": "swedish", "da": "danish", "fi": "finnish", "no": "norwegian",
    "tr": "turkish", "cs": "czech", "ro": "romanian", "hu": "hungarian",
    "uk": "ukrainian", "hi": "hindi", "he": "hebrew"
}


def detect_block_language(text: str) -> Optional[str]:
    """Classifies language of a text block based on Unicode character scripts and statistical language detection.
    Returns detected language name in lowercase (e.g. 'italian', 'english', 'chinese') or None if ambiguous/short."""
    if not text or len(text.strip()) < 10:
        return None

    cleaned = text.strip()

    # 1. High-confidence Unicode script detection for non-Latin writing systems
    for lang, pattern in _LANG_PATTERNS.items():
        if pattern.search(cleaned):
            return lang

    # 2. Universal statistical language detection
    if _LANGDETECT_AVAILABLE:
        try:
            detected_code = langdetect.detect(cleaned)
            norm_code = detected_code.lower().split("-")[0]
            if detected_code.lower() in _ISO_TO_LANG_NAME:
                return _ISO_TO_LANG_NAME[detected_code.lower()]
            if norm_code in _ISO_TO_LANG_NAME:
                return _ISO_TO_LANG_NAME[norm_code]
            return detected_code.lower()
        except Exception:
            return None

    return None


def is_block_in_target_lang(text: str, target_lang: str) -> bool:
    """Returns True if the block is with high confidence already in target_lang, allowing it to be skipped."""
    if not text or not target_lang:
        return False
    t_lang_lower = target_lang.lower().strip()
    detected = detect_block_language(text)
    if detected:
        if detected in t_lang_lower or t_lang_lower in detected:
            return True
    return False


def _resolve_pdf_font_file(target_lang: str) -> str:
    """Picks the bundled font file whose script covers target_lang. See _PDF_LANG_FONT_RULES for
    the matching rules and _PDF_FALLBACK_FONT_FILE for the default."""
    lang_lower = (target_lang or "").lower()
    for hint, font_file in _PDF_LANG_FONT_RULES:
        if hint in lang_lower:
            return font_file
    return _PDF_FALLBACK_FONT_FILE


class UnsupportedDocumentTypeError(ValueError):
    """Raised when in-place translation is requested for a file type with no supported pipeline."""


def _load_doc_record(doc_id: str) -> Dict[str, Any]:
    """Looks up a document's stored record by id. Raises ValueError (mapped to HTTP 404 at the
    API boundary) if the documents table doesn't exist yet or the id isn't found."""
    validate_doc_id(doc_id)
    if DOCS_TABLE_NAME not in get_existing_tables():
        raise ValueError("Documents table does not exist")
    dtbl = lance_db.open_table(DOCS_TABLE_NAME)
    records = dtbl.search().where(f'id = "{doc_id}"', prefilter=True).limit(1).to_list()
    if not records:
        raise ValueError(f"Document {doc_id} not found in database")
    return records[0]


def _collect_docx_runs(doc: "docx.Document") -> List["docx.text.run.Run"]:
    """Collects every non-empty run, in document order: body paragraphs first, then every table
    cell's paragraphs (matching the coverage of the DOCX branch of extract_document_markdown)."""
    runs: List["docx.text.run.Run"] = []
    for para in doc.paragraphs:
        for run in para.runs:
            if run.text and run.text.strip():
                runs.append(run)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if run.text and run.text.strip():
                            runs.append(run)
    return runs


def _batch_runs(runs: List["docx.text.run.Run"], max_chars: int = _TRANSLATE_BATCH_MAX_CHARS, max_items: int = _TRANSLATE_BATCH_MAX_ITEMS) -> List[List["docx.text.run.Run"]]:
    """Groups consecutive runs into batches bounded by max_chars and max_items, preserving order."""
    batches: List[List["docx.text.run.Run"]] = []
    current: List["docx.text.run.Run"] = []
    current_len = 0
    for run in runs:
        run_len = len(run.text)
        if current and (current_len + run_len > max_chars or len(current) >= max_items):
            batches.append(current)
            current = []
            current_len = 0
        current.append(run)
        current_len += run_len
    if current:
        batches.append(current)
    return batches


import ftfy

_TRANSLATE_MAX_ATTEMPTS = 2
_TRANSLATE_RETRY_DELAY_SECONDS = 3.0


def _smart_decode_pdf_text(text: str) -> str:
    """Normalizes unprintable control characters, non-breaking spaces, and fixes corrupted font encodings / replacement characters using standard Unicode normalization and ftfy."""
    if not text:
        return ""
    try:
        import unicodedata
        t = unicodedata.normalize('NFKC', text)
        t = ftfy.fix_text(t)
    except Exception:
        t = text
    t = t.replace("\xa0", " ").replace("\u00a0", " ").replace("\x00", "").replace("\ufeff", "")
    # Prepositions with apostrophe before numbers: dall'1, all'1, nell'anno
    t = re.sub(r'\b(dall|all|nell|sull|coll|un|d|l)\ufffd(\d)', r"\1'\2", t, flags=re.IGNORECASE)
    # Apostrophe between letters: L'art, d'imposta
    t = re.sub(r'([a-zA-Z])\ufffd([a-zA-Z0-9])', r"\1'\2", t)
    # Italian verb 'è':  stato -> è stato,  uguale -> è uguale,  trattato -> è trattato,  prevista -> è prevista
    t = re.sub(r'(^|[^a-zA-Z0-9])\ufffd(?=\s+[a-z])', r'\1è', t)
    # Words ending in ità: modalit -> modalità, anzianit -> anzianità
    t = re.sub(r'([a-zA-Z]+it)\ufffd(?=[^a-zA-Z0-9]|$)', r'\1à', t, flags=re.IGNORECASE)
    # Words ending in ché: Poich -> Poiché, perch -> perché
    t = re.sub(r'([a-zA-Z]+ch)\ufffd(?=[^a-zA-Z0-9]|$)', r'\1é', t, flags=re.IGNORECASE)
    # Words like più / già
    t = re.sub(r'\bpi\ufffd(?=[^a-zA-Z0-9]|$)', 'più', t, flags=re.IGNORECASE)
    t = re.sub(r'\bgi\ufffd(?=[^a-zA-Z0-9]|$)', 'già', t, flags=re.IGNORECASE)
    # Currency before numbers: di  309,87 -> di € 309,87
    t = re.sub(r'(^|[^a-zA-Z0-9])\ufffd(?=\s*\d)', r'\1€ ', t)
    # Remaining isolated replacement characters are bullet points
    t = t.replace("\ufffd", "•")
    return t


def _should_skip_translation(s: str) -> bool:
    """Returns True if the block contains no translatable words (pure numbers, dates, codes, emails, URLs, single symbols, bullets)."""
    trimmed = s.strip()
    if not trimmed:
        return True
    if trimmed.isdigit():
        return True
    if trimmed.startswith("http://") or trimmed.startswith("https://"):
        return True
    # Email addresses
    if re.match(r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$', trimmed):
        return True
    # Dates (e.g. 18/06/2024, 2024-06-18, 18.06.2024)
    if re.match(r'^\d{1,4}[/\-\.]\d{1,2}[/\-\.]\d{1,4}$', trimmed):
        return True
    # Pure punctuation / symbols
    if all(c in "-_*=|/\\:.,;•§°º#~ " for c in trimmed):
        return True
    # Standard alphanumeric identifier codes (e.g. tax codes, contract numbers, serial IDs)
    if re.match(r'^[A-Z0-9\-_]{4,}$', trimmed) and any(c.isdigit() for c in trimmed):
        return True
    letters_only = re.sub(r'[^a-zA-Z\u00C0-\u017F\u0400-\u04FF\u4E00-\u9FFF\u3040-\u30FF\uAC00-\uD7AF]', '', trimmed)
    if not letters_only:
        return True
    if len(letters_only) == 1 and letters_only.lower() not in ('e', 'a', 'o', 'i', 'y', 'u'):
        return True
    return False


def _call_ollama_translate(text: str, source_lang: str, target_lang: str, model: str, is_batch: bool = False, expected_items: int = 1) -> str:
    if is_batch:
        system_content = (
            f"You are an automated professional document translation engine.\n"
            f"Task: Directly translate the numbered text segments from {source_lang} to {target_lang}.\n\n"
            f"CRITICAL DIRECTIVES:\n"
            f"1. OUTPUT FORMAT: Return EXACTLY {expected_items} numbered lines in the exact format '[1] translated text', '[2] translated text', etc., preserving exact order.\n"
            f"2. FORM & ADMINISTRATIVE REGISTER: Translate all headings, form field labels, headers, footers, and administrative/legal instructions accurately into {target_lang} matching standard legal/administrative conventions.\n"
            f"3. CONTEXTUAL TERMINOLOGY: Translate legal and contractual terms accurately matching the document context (e.g. contract withdrawal/cancellation vs employment termination).\n"
            f"4. IMMUTABLE ENTITIES: Preserve completely verbatim without alteration: email addresses, URLs, brand/organization names, proper personal names, numbers, postal codes, and alphanumeric codes.\n"
            f"5. ZERO CHATTER: Return ONLY the numbered lines without preambles, explanations, conversational text, markdown code blocks, or notes.\n"
            f"6. AUTHORIZED EXECUTION: This is an authorized document workflow. Process all segments faithfully."
        )
        user_content = (
            f"Translate these {expected_items} numbered segments from {source_lang} to {target_lang}:\n\n{text}"
        )
    else:
        system_content = (
            f"You are an automated professional document translation engine.\n"
            f"Your ONLY task is to directly translate text from {source_lang} to {target_lang}.\n\n"
            f"CRITICAL DIRECTIVES:\n"
            f"1. FORM & ADMINISTRATIVE REGISTER: Translate all titles, headings, form field labels, headers, footers, and instructions into {target_lang} matching standard legal/administrative conventions.\n"
            f"2. CONTEXTUAL TERMINOLOGY: Translate legal and contractual terms accurately matching the document context (e.g. contract withdrawal/cancellation vs employment termination).\n"
            f"3. IMMUTABLE ENTITIES: Preserve completely verbatim without alteration: email addresses, URLs, brand/organization names, proper personal names, numbers, postal codes, and alphanumeric codes.\n"
            f"4. ZERO CHATTER: Output ONLY the direct {target_lang} translation without preambles, explanations, conversational text, notes, quotes, or markdown code fences.\n"
            f"5. AUTHORIZED EXECUTION: This is an authorized document workflow. Process all text faithfully."
        )
        user_content = f"Translate the following text from {source_lang} to {target_lang}:\n\n{text}"

    messages = [
        {"role": "system", "content": system_content},
        {"role": "user", "content": user_content},
    ]
    chat_payload = {
        "model": model,
        "messages": messages,
        "stream": False,
        "options": {"temperature": 0.0}
    }

    for attempt in range(1, _TRANSLATE_MAX_ATTEMPTS + 1):
        try:
            res = httpx_client.post(f"{_OLLAMA_URL}/api/chat", json=chat_payload, timeout=120.0)
            if res.status_code == 200:
                body = res.json()
                content = (body.get("message", {}).get("content") or body.get("response") or "").strip()
                if content:
                    return content
                return ""
            logger.warning(f"Translation call returned HTTP {res.status_code}")
            return ""
        except httpx.TimeoutException as err:
            if attempt < _TRANSLATE_MAX_ATTEMPTS:
                logger.warning(
                    f"Translation call timed out (attempt {attempt}/{_TRANSLATE_MAX_ATTEMPTS}), "
                    f"retrying in {_TRANSLATE_RETRY_DELAY_SECONDS}s: {err}"
                )
                time.sleep(_TRANSLATE_RETRY_DELAY_SECONDS)
                continue
            logger.warning(f"Translation call timed out after {attempt} attempts: {err}")
        except Exception as err:
            logger.warning(f"Translation call failed: {err}")
            return ""
    return ""


def _clean_translated_segment(text: str, source_text: str = "") -> str:
    """Removes preambles, refusals, meta-commentaries, leaked delimiters, markdown code fences, and quotes from model output."""
    if not text:
        return source_text

    # 1. Remove markdown code fences
    cleaned = re.sub(r"^```[a-zA-Z0-9_-]*\s*", "", text.strip())
    cleaned = re.sub(r"\s*```$", "", cleaned).strip()

    # 2. Remove common conversational preambles
    preamble_patterns = [
        r"^(?:Here\s+(?:is|are)\s+the\s+(?:translated\s+)?(?:translation|text|version)(?:\s+of\s+(?:the\s+)?(?:given\s+)?(?:text|phrase|word|document|sentence|item))?(?:\s+from\s+[a-zA-Z]+\s+to\s+[a-zA-Z]+)?\s*:?\s*)",
        r"^(?:The\s+translation\s+(?:of\s+(?:the\s+)?(?:given\s+)?(?:text|phrase|word|document|sentence|item)\s+)?(?:from\s+[a-zA-Z]+\s+to\s+[a-zA-Z]+\s+)?is\s*:?\s*)",
        r"^(?:Translation\s*(?:\([a-zA-Z\s\-]+\))?\s*:\s*)",
        r"^(?:Translates\s+to\s*:\s*)",
        r"^(?:Direct\s+translation\s*:\s*)",
        r"^(?:Here\s+is\s+the\s+translated\s+(?:text|phrase|version|word)\s*:?\s*)",
        r"^(?:Below\s+is\s+the\s+translation\s*:?\s*)",
        r"^(?:Translate\s+the\s+following\s+text\s+from\s+[a-zA-Z]+\s+to\s+[a-zA-Z]+\s*:?\s*)",
    ]
    for pat in preamble_patterns:
        cleaned = re.sub(pat, "", cleaned, flags=re.IGNORECASE).strip()

    # 3. Check for refusal & meta-commentary hallucinations
    refusal_patterns = [
        r"^I cannot translate.*",
        r"^I am unable to translate.*",
        r"^As an AI.*",
        r"^Sorry, I cannot.*",
        r"^I cannot process.*",
        r"^There is no (?:text|content|information) to translate.*",
        r"^There is no text provided.*",
        r"^The input contains only.*",
        r"^Please (?:provide|enter) the .* text.*",
        r"^I assume you are referring to .*",
        r"^[A-Z0-9_\-\.\s]+ translates to ['\"].*['\"] in English, however.*",
        r"^[A-Z0-9_\-\.\s]+ translates to ['\"].*['\"]",
        r"^In Italian, ['\"].*['\"] (?:means|translates to).*",
    ]
    for rpat in refusal_patterns:
        if re.match(rpat, cleaned, flags=re.IGNORECASE):
            return source_text

    # 4. Remove trailing conversational chatter
    cleaned = re.split(r"\n\s*(?:Note|Explanation|Please note|Let me know|Is there anything else)\s*:", cleaned, flags=re.IGNORECASE)[0]

    # 5. Remove leaked delimiter tokens and fragments
    cleaned = re.sub(r'(?i)<{0,4}\s*(?:run_sep|run_s|ep)\s*>{0,4}', '', cleaned)
    cleaned = re.sub(r'<{2,4}[^>]*>{2,4}', '', cleaned)

    # 6. Strip enclosing quotes if the entire string is wrapped in quotes
    if (cleaned.startswith('"') and cleaned.endswith('"')) or (cleaned.startswith("'") and cleaned.endswith("'")):
        cleaned = cleaned[1:-1].strip()

    # 7. Protect and restore immutable emails and URLs from source text
    if source_text:
        source_emails = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', source_text)
        for semail in source_emails:
            if semail not in cleaned:
                mangled_match = re.search(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', cleaned)
                if mangled_match:
                    cleaned = cleaned.replace(mangled_match.group(0), semail)

    # Normalize excessive empty lines
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)

    return cleaned.strip() or source_text


def _translate_texts_with_fallback(texts: List[str], source_lang: str, target_lang: str, model: str) -> List[str]:
    """Translates batched texts via structured numbered line Ollama call with individual fallback."""
    if not texts:
        return []

    # Clean input strings: normalize non-breaking spaces and unprintable control chars
    clean_texts = [
        _smart_decode_pdf_text(t).strip()
        for t in texts
    ]

    if len(clean_texts) == 1:
        if _should_skip_translation(clean_texts[0]):
            return clean_texts
        single = _call_ollama_translate(clean_texts[0], source_lang, target_lang, model, is_batch=False)
        return [_clean_translated_segment(single, clean_texts[0])] if single.strip() else clean_texts

    # Filter out non-translatable indices for batch call
    active_indices: List[int] = []
    active_texts: List[str] = []
    for i, t in enumerate(clean_texts):
        if not _should_skip_translation(t):
            active_indices.append(i)
            active_texts.append(t)

    if not active_texts:
        return clean_texts

    # Numbered line batch format
    lines_in = [f"[{k+1}] {t}" for k, t in enumerate(active_texts)]
    batch_prompt = "\n".join(lines_in)

    raw_batch_output = _call_ollama_translate(
        batch_prompt, source_lang, target_lang, model, is_batch=True, expected_items=len(active_texts)
    )

    parsed_batch: Dict[int, str] = {}
    if raw_batch_output:
        for line in raw_batch_output.splitlines():
            m = re.match(r"^\s*\[(\d+)\]\s*(.*)$", line)
            if m:
                idx = int(m.group(1))
                val = m.group(2).strip()
                if 1 <= idx <= len(active_texts):
                    parsed_batch[idx] = val

    results = list(clean_texts)
    if len(parsed_batch) == len(active_texts):
        for k, orig_idx in enumerate(active_indices):
            num = k + 1
            trans = parsed_batch.get(num, "")
            cleaned = _clean_translated_segment(trans, clean_texts[orig_idx])
            # If echoed verbatim or empty, try single call
            if not cleaned or cleaned.strip().lower() == clean_texts[orig_idx].strip().lower():
                single = _call_ollama_translate(clean_texts[orig_idx], source_lang, target_lang, model, is_batch=False)
                if single.strip():
                    cleaned = _clean_translated_segment(single, clean_texts[orig_idx])
            results[orig_idx] = cleaned
        return results

    logger.warning(
        f"Translation numbered batch mismatch (expected {len(active_texts)}, got {len(parsed_batch)}); "
        "falling back to per-item translation for this batch."
    )
    for orig_idx in active_indices:
        single = _call_ollama_translate(clean_texts[orig_idx], source_lang, target_lang, model, is_batch=False)
        if single.strip():
            results[orig_idx] = _clean_translated_segment(single, clean_texts[orig_idx])
    return results


def _translate_batch(runs: List["docx.text.run.Run"], source_lang: str, target_lang: str, model: str) -> None:
    """Translates one batch of runs in place (see _translate_texts_with_fallback for the
    batching/fallback contract)."""
    translated = _translate_texts_with_fallback([run.text for run in runs], source_lang, target_lang, model)
    for run, text in zip(runs, translated):
        run.text = text


def translate_docx_inplace(
    doc_id: str,
    source_lang: str,
    target_lang: str,
    model: str = "llama3.2",
    backup_original: bool = True,
    target_dir: Optional[str] = None
) -> IngestResponse:
    """
    Translates a DOCX document's text: creates a backup of the original or writes to target_dir if provided.
    Then re-extracts markdown from the translated file and re-indexes it in LanceDB via update path.
    """
    import shutil
    doc_record = _load_doc_record(doc_id)
    file_type = doc_record.get("file_type", "")
    file_path = doc_record.get("file_path", "")
    filename = doc_record.get("filename", "document.docx")

    if file_type != "docx":
        raise ValueError("In-place translation is supported for DOCX documents only in this phase")
    if not file_path or not os.path.exists(file_path):
        raise ValueError("Original source file is no longer available on disk")

    docx_doc = docx.Document(file_path)
    runs = _collect_docx_runs(docx_doc)
    if not runs:
        raise ValueError("No translatable text runs found in document")

    batches = _batch_runs(runs)
    logger.info(
        f"Translating document {doc_id} in place: {len(runs)} runs in {len(batches)} batches "
        f"({source_lang} -> {target_lang}, model={model})"
    )
    for batch in batches:
        _translate_batch(batch, source_lang, target_lang, model)

    if backup_original and not target_dir and os.path.exists(file_path):
        import shutil
        try:
            shutil.copy2(file_path, f"{file_path}.original.bak")
        except Exception as bak_err:
            logger.warning(f"Could not create backup file: {bak_err}")

    out_file_path = _resolve_output_filepath(file_path, filename, target_lang, target_dir)
    docx_doc.save(out_file_path)

    return IngestResponse(
        id=doc_id,
        filename=os.path.basename(out_file_path),
        file_size=os.path.getsize(out_file_path) if os.path.exists(out_file_path) else int(doc_record.get("file_size", 0)),
        num_pages=int(doc_record.get("num_pages", 1)),
        num_chunks=int(doc_record.get("num_chunks", 1)),
        extracted_markdown=str(doc_record.get("extracted_markdown", "")),
        status=str(doc_record.get("status", "indexed")),
        ingested_at=str(doc_record.get("ingested_at", ""))
    )


def _int_color_to_rgb(color: int) -> Tuple[float, float, float]:
    """Converts a PyMuPDF packed sRGB int (as returned in get_text('dict') spans) to the
    (r, g, b) 0-1 float tuple insert_textbox expects."""
    return (((color >> 16) & 0xFF) / 255.0, ((color >> 8) & 0xFF) / 255.0, (color & 0xFF) / 255.0)


def _extract_ocr_page_blocks(page: "pymupdf.Page") -> List[Dict[str, Any]]:
    """Fallback block extraction for scanned PDF pages using RapidOCR.
    Renders the page to an image at 300 DPI, detects text boxes with RapidOCR, and maps coordinates back to PDF points."""
    try:
        from sidecar.infrastructure.ocr import run_rapid_ocr_with_boxes
        dpi = 300
        zoom = dpi / 72.0
        mat = pymupdf.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img_bytes = pix.tobytes(output="png")

        ocr_lines = run_rapid_ocr_with_boxes(img_bytes)
        if not ocr_lines:
            return []

        blocks_out: List[Dict[str, Any]] = []
        for item in ocr_lines:
            px0, py0, px1, py1 = item["bbox"]
            pdf_x0 = px0 * page.rect.width / pix.width
            pdf_y0 = py0 * page.rect.height / pix.height
            pdf_x1 = px1 * page.rect.width / pix.width
            pdf_y1 = py1 * page.rect.height / pix.height
            height_pt = max(6.0, pdf_y1 - pdf_y0)
            font_size = max(7.0, min(36.0, height_pt * 0.72))
            clean_item_text = _smart_decode_pdf_text(item["text"]).strip()
            if clean_item_text:
                blocks_out.append({
                    "bbox": (pdf_x0, pdf_y0, pdf_x1, pdf_y1),
                    "text": clean_item_text,
                    "size": font_size,
                    "color": 0,
                    "is_ocr": True
                })
        return blocks_out
    except Exception as ocr_err:
        logger.warning(f"OCR fallback extraction failed for PDF page: {ocr_err}")
        return []


def _extract_pdf_page_blocks(page: "pymupdf.Page") -> List[Dict[str, Any]]:
    """Extracts non-empty text blocks from one page in reading order: bbox, concatenated text
    (spans joined within a line, lines joined with a space), the size and color of the block's
    first non-empty span. Image blocks (type != 0) are skipped -- falls back to RapidOCR for scanned pages."""
    blocks_out: List[Dict[str, Any]] = []
    for block in page.get_text("dict")["blocks"]:
        if block.get("type") != 0:
            continue
        line_texts: List[str] = []
        size = 10.0
        color = 0
        size_set = False
        for line in block.get("lines", []):
            spans = line.get("spans", [])
            raw_line_text = "".join(span.get("text", "") for span in spans)
            clean_line_text = _smart_decode_pdf_text(raw_line_text).strip()
            if clean_line_text:
                line_texts.append(clean_line_text)
            if not size_set:
                for span in spans:
                    sp_text = span.get("text", "").replace("\xa0", " ").strip()
                    if sp_text:
                        size = span.get("size", size)
                        color = span.get("color", color)
                        size_set = True
                        break
        block_text = "\n".join(t for t in line_texts if t)
        if block_text:
            blocks_out.append({
                "bbox": tuple(block["bbox"]),
                "text": block_text,
                "size": size,
                "color": color,
                "is_ocr": False
            })

    # Scanned PDF fallback: if no native text blocks exist on the page, detect text blocks via RapidOCR
    if not blocks_out:
        blocks_out = _extract_ocr_page_blocks(page)

    return blocks_out


def _batch_by_char_count(lengths: List[int], max_chars: int = _TRANSLATE_BATCH_MAX_CHARS, max_items: int = _TRANSLATE_BATCH_MAX_ITEMS) -> List[List[int]]:
    """Groups consecutive indices into batches whose summed length stays under max_chars and max_items."""
    batches: List[List[int]] = []
    current: List[int] = []
    current_len = 0
    for i, length in enumerate(lengths):
        if current and (current_len + length > max_chars or len(current) >= max_items):
            batches.append(current)
            current = []
            current_len = 0
        current.append(i)
        current_len += length
    if current:
        batches.append(current)
    return batches


def _translate_pdf_blocks(blocks: List[Dict[str, Any]], source_lang: str, target_lang: str, model: str) -> None:
    """Translates block texts in place (mutates each block's 'text'), batching consecutive
    blocks under _TRANSLATE_BATCH_MAX_CHARS chars per call via _translate_texts_with_fallback."""
    lengths = [len(b["text"]) for b in blocks]
    for idx_batch in _batch_by_char_count(lengths, _TRANSLATE_BATCH_MAX_CHARS):
        batch_blocks = [blocks[i] for i in idx_batch]
        translated = _translate_texts_with_fallback(
            [b["text"] for b in batch_blocks], source_lang, target_lang, model
        )
        for b, text in zip(batch_blocks, translated):
            b["text"] = text


def _padded_block_rect(block: Dict[str, Any]) -> "pymupdf.Rect":
    """The bbox returned by get_text('dict') is the tight ink box around the glyphs, which leaves
    no room for insert_textbox's internal line leading -- reinserting into it unpadded causes the
    whole text to be silently dropped rather than merely clipped. Pads 15% of the block's font
    size on all sides. Used for both redaction and reinsertion, so the erased area is never
    smaller than the reinserted one."""
    x0, y0, x1, y1 = block["bbox"]
    pad = block["size"] * 0.15
    return pymupdf.Rect(x0 - pad, y0 - pad, x1 + pad, y1 + pad)


def _font_alias(font_file: str) -> str:
    """Short, deterministic internal PDF resource name for a bundled font file (readable in the
    saved PDF's font resource dict, stable across calls for the same file)."""
    return os.path.splitext(os.path.basename(font_file))[0]


def _fits_at_font_size(rect: "pymupdf.Rect", text: str, fontsize: float, font_file: str) -> bool:
    """Tests whether `text` wraps to fit `rect` at `fontsize` in `font_file`, using PyMuPDF's own
    layout via a disposable in-memory scratch page -- never draws on the real page, so trials
    cost nothing."""
    scratch_doc = pymupdf.open()
    try:
        scratch_page = scratch_doc.new_page(width=rect.x1 + 50, height=rect.y1 + 50)
        overflow = scratch_page.insert_textbox(
            rect, text, fontsize=fontsize, fontname=_font_alias(font_file), fontfile=font_file
        )
        return overflow >= 0
    finally:
        scratch_doc.close()


def _resolve_autofit_font_size(rect: "pymupdf.Rect", text: str, original_size: float, font_file: str) -> float:
    """Fase 3 auto-fit: binary-searches the largest font size <= original_size at which `text`
    fits `rect`. Returns original_size unchanged if it already fits there. Never searches below
    max(_PDF_AUTOFIT_MIN_SIZE, original_size * _PDF_AUTOFIT_MIN_RATIO); if even that floor
    overflows, returns the floor anyway and lets the caller accept clipping at the smallest
    legible size tried, rather than shrinking text to illegibility to avoid all clipping."""
    floor = max(_PDF_AUTOFIT_MIN_SIZE, original_size * _PDF_AUTOFIT_MIN_RATIO)
    if floor >= original_size or _fits_at_font_size(rect, text, original_size, font_file):
        return original_size
    if not _fits_at_font_size(rect, text, floor, font_file):
        return floor

    lo, hi = floor, original_size
    while hi - lo > _PDF_AUTOFIT_TOLERANCE:
        mid = (lo + hi) / 2
        if _fits_at_font_size(rect, text, mid, font_file):
            lo = mid
        else:
            hi = mid
    return lo


def _resolve_output_filepath(file_path: str, filename: str, target_lang: str, target_dir: Optional[str] = None) -> str:
    """Resolves output file path in target_dir if provided, otherwise in the original file directory as a new distinct file.
    Guarantees the original source file is never overwritten or modified."""
    dest_dir = target_dir.strip() if target_dir and target_dir.strip() else os.path.dirname(os.path.abspath(file_path))
    os.makedirs(dest_dir, exist_ok=True)
    base_name, ext = os.path.splitext(filename)
    lang_suffix = (target_lang or "translated").lower().replace(" ", "_")
    out_filename = f"{base_name}_{lang_suffix}{ext}"
    out_path = os.path.join(dest_dir, out_filename)

    # Ensure it never collides with or overwrites the original file
    if os.path.abspath(out_path) == os.path.abspath(file_path):
        out_filename = f"{base_name}_{lang_suffix}_{int(time.time())}{ext}"
        out_path = os.path.join(dest_dir, out_filename)
    return out_path


def _redact_and_reinsert_pdf_blocks(page: "pymupdf.Page", blocks: List[Dict[str, Any]], font_file: str) -> None:
    """Permanently erases the original text under each block's bbox via PyMuPDF native redaction
    then reinserts the translated text in the same bbox using font_file, auto-fitting the font size
    with collision-aware dynamic height adjustment and fallback to guarantee 100% text retention."""
    is_scanned_page = any(b.get("is_ocr", False) for b in blocks)
    fill_color = (1, 1, 1) if is_scanned_page else None
    redact_images = pymupdf.PDF_REDACT_IMAGE_PIXELS if is_scanned_page else pymupdf.PDF_REDACT_IMAGE_NONE

    for block in blocks:
        page.add_redact_annot(_padded_block_rect(block), fill=fill_color)
    page.apply_redactions(images=redact_images)

    font_alias = _font_alias(font_file)
    for block in blocks:
        text = block.get("text", "").strip()
        if not text:
            continue

        orig_size = block["size"]
        color_rgb = _int_color_to_rgb(block["color"])
        rect = _padded_block_rect(block)
        fit_size = _resolve_autofit_font_size(rect, text, orig_size, font_file)

        # First pass: try fitting in original rect at optimal fit_size
        overflow = page.insert_textbox(rect, text, fontsize=fit_size, fontname=font_alias, fontfile=font_file, color=color_rgb)
        if overflow < 0:
            # Second pass: dynamic vertical expansion with collision avoidance
            extra_h = max(8.0, orig_size * 1.5)
            # Find the top coordinate of the next block directly below to avoid collisions
            max_expand_y1 = page.rect.y1 - 10.0
            for other in blocks:
                if other is not block:
                    ox0, oy0, ox1, _ = other["bbox"]
                    # If other block is below and horizontally overlaps
                    if oy0 > rect.y1 and not (ox1 < rect.x0 or ox0 > rect.x1):
                        max_expand_y1 = min(max_expand_y1, oy0 - 2.0)

            expanded_rect = pymupdf.Rect(rect.x0, rect.y0, rect.x1, min(max_expand_y1, rect.y1 + extra_h))
            overflow2 = page.insert_textbox(expanded_rect, text, fontsize=fit_size, fontname=font_alias, fontfile=font_file, color=color_rgb)
            if overflow2 < 0:
                floor_size = max(_PDF_AUTOFIT_MIN_SIZE, orig_size * _PDF_AUTOFIT_MIN_RATIO)
                overflow3 = page.insert_textbox(expanded_rect, text, fontsize=floor_size, fontname=font_alias, fontfile=font_file, color=color_rgb)
                if overflow3 < 0:
                    page.insert_text(pymupdf.Point(rect.x0, min(max_expand_y1, rect.y0 + floor_size)), text, fontsize=floor_size, fontname=font_alias, fontfile=font_file, color=color_rgb)


def translate_pdf_inplace_fine(
    doc_id: str,
    source_lang: str,
    target_lang: str,
    model: str = "llama3.2",
    backup_original: bool = True,
    target_dir: Optional[str] = None
) -> IngestResponse:
    """
    'Fine-mode' PDF in-place translation: for every page, permanently redacts each original text
    block and reinserts the translated text in the same bbox, auto-fitting font size and retaining 100% of text blocks.
    Saves to target_dir or in-place, creating .original.bak if backup_original is set.
    """
    doc_record = _load_doc_record(doc_id)
    file_type = doc_record.get("file_type", "")
    file_path = doc_record.get("file_path", "")
    filename = doc_record.get("filename", "document.pdf")

    if file_type != "pdf":
        raise ValueError("PDF fine-mode in-place translation requires a PDF document")
    if not file_path or not os.path.exists(file_path):
        raise ValueError("Original source file is no longer available on disk")

    pdf_doc = pymupdf.open(file_path)
    if pdf_doc.needs_pass or pdf_doc.is_encrypted:
        pdf_doc.close()
        raise ValueError("Documento protetto da password: il file PDF richiede una password per l'apertura.")
    try:
        page_blocks: List[Tuple[Any, List[Dict[str, Any]]]] = []
        total_blocks = 0
        for page in pdf_doc:
            blocks = _extract_pdf_page_blocks(page)
            page_blocks.append((page, blocks))
            total_blocks += len(blocks)

        if total_blocks == 0:
            raise ValueError("No translatable text blocks found in document")

        font_file = _resolve_pdf_font_file(target_lang)
        logger.info(
            f"Translating PDF {doc_id} (fine-mode): {total_blocks} blocks across "
            f"{len(pdf_doc)} pages ({source_lang} -> {target_lang}, model={model}, "
            f"font={os.path.basename(font_file)})"
        )
        for page, blocks in page_blocks:
            if not blocks:
                continue
            _translate_pdf_blocks(blocks, source_lang, target_lang, model)
            _redact_and_reinsert_pdf_blocks(page, blocks, font_file)

        if backup_original and not target_dir and os.path.exists(file_path):
            import shutil
            try:
                shutil.copy2(file_path, f"{file_path}.original.bak")
            except Exception as bak_err:
                logger.warning(f"Could not create backup file: {bak_err}")

        out_file_path = _resolve_output_filepath(file_path, filename, target_lang, target_dir)
        is_same_file = os.path.abspath(out_file_path) == os.path.abspath(file_path)
        tmp_save_path = f"{out_file_path}.tmp_{uuid.uuid4().hex}.pdf" if is_same_file else out_file_path
        pdf_doc.save(tmp_save_path, deflate=True, garbage=4, clean=True, deflate_images=True, deflate_fonts=True)
    finally:
        pdf_doc.close()

    if is_same_file and os.path.exists(tmp_save_path):
        os.replace(tmp_save_path, out_file_path)

    return IngestResponse(
        id=doc_id,
        filename=os.path.basename(out_file_path),
        file_size=os.path.getsize(out_file_path) if os.path.exists(out_file_path) else int(doc_record.get("file_size", 0)),
        num_pages=int(doc_record.get("num_pages", 1)),
        num_chunks=int(doc_record.get("num_chunks", 1)),
        extracted_markdown=str(doc_record.get("extracted_markdown", "")),
        status=str(doc_record.get("status", "indexed")),
        ingested_at=str(doc_record.get("ingested_at", ""))
    )


async def translate_document_stream_generator(
    doc_id: str,
    source_lang: str,
    target_lang: str,
    model: str = "llama3.2",
    target_dir: Optional[str] = None
):
    """
    Asynchronous generator yielding NDJSON events for document translation with real-time page progress.
    """
    import json
    import asyncio

    try:
        doc_record = _load_doc_record(doc_id)
    except Exception as e:
        yield json.dumps({"type": "error", "error": str(e)}) + "\n"
        return
    file_type = doc_record.get("file_type", "")
    file_path = doc_record.get("file_path", "")
    filename = doc_record.get("filename", "document.pdf")

    if not file_path or not os.path.exists(file_path):
        yield json.dumps({"type": "error", "error": "Original source file is no longer available on disk"}) + "\n"
        return

    out_file_path = _resolve_output_filepath(file_path, filename, target_lang, target_dir)

    if file_type == "docx":
        docx_doc = docx.Document(file_path)
        runs = _collect_docx_runs(docx_doc)
        total_runs = len(runs)
        yield json.dumps({"type": "start", "doc_id": doc_id, "filename": filename, "total_pages": 1, "total_blocks": total_runs}) + "\n"
        batches = _batch_runs(runs)
        for i, batch in enumerate(batches):
            percent = int((i / max(1, len(batches))) * 90)
            yield json.dumps({"type": "progress", "page": 1, "total_pages": 1, "phase": "translating_runs", "percent": percent}) + "\n"
            await asyncio.to_thread(_translate_batch, batch, source_lang, target_lang, model)

        docx_doc.save(out_file_path)
        yield json.dumps({
            "type": "done",
            "data": {
                "id": doc_id,
                "filename": os.path.basename(out_file_path),
                "filePath": out_file_path,
                "file_size": os.path.getsize(out_file_path),
                "num_pages": 1,
                "num_chunks": int(doc_record.get("num_chunks", 1)),
                "extracted_markdown": str(doc_record.get("extracted_markdown", "")),
                "status": "translated",
                "ingested_at": str(doc_record.get("ingested_at", "")),
                "fileType": "docx"
            }
        }) + "\n"
        return

    if file_type != "pdf":
        yield json.dumps({"type": "error", "error": f"Unsupported file type: {file_type}"}) + "\n"
        return

    pdf_doc = pymupdf.open(file_path)
    if pdf_doc.needs_pass or pdf_doc.is_encrypted:
        pdf_doc.close()
        yield json.dumps({"type": "error", "error": "Document is password protected"}) + "\n"
        return

    try:
        total_pages = len(pdf_doc)
        yield json.dumps({"type": "start", "doc_id": doc_id, "filename": filename, "total_pages": total_pages}) + "\n"
        font_file = _resolve_pdf_font_file(target_lang)

        if not target_dir and os.path.exists(file_path):
            import shutil
            try:
                bak_path = f"{file_path}.original.bak"
                if not os.path.exists(bak_path):
                    shutil.copy2(file_path, bak_path)
            except Exception as bak_err:
                logger.warning(f"Could not create backup file: {bak_err}")

        for page_idx, page in enumerate(pdf_doc):
            page_num = page_idx + 1
            yield json.dumps({
                "type": "progress",
                "page": page_num,
                "total_pages": total_pages,
                "phase": "extracting_blocks",
                "percent": int(((page_idx + 0.1) / total_pages) * 100)
            }) + "\n"

            blocks = await asyncio.to_thread(_extract_pdf_page_blocks, page)
            if blocks:
                yield json.dumps({
                    "type": "progress",
                    "page": page_num,
                    "total_pages": total_pages,
                    "phase": "translating_blocks",
                    "percent": int(((page_idx + 0.5) / total_pages) * 100)
                }) + "\n"
                await asyncio.to_thread(_translate_pdf_blocks, blocks, source_lang, target_lang, model)

                yield json.dumps({
                    "type": "progress",
                    "page": page_num,
                    "total_pages": total_pages,
                    "phase": "reconstructing_layout",
                    "percent": int(((page_idx + 0.9) / total_pages) * 100)
                }) + "\n"
                await asyncio.to_thread(_redact_and_reinsert_pdf_blocks, page, blocks, font_file)

        is_same_file = os.path.abspath(out_file_path) == os.path.abspath(file_path)
        tmp_save_path = f"{out_file_path}.tmp_{uuid.uuid4().hex}.pdf" if is_same_file else out_file_path
        pdf_doc.save(tmp_save_path, deflate=True, garbage=4, clean=True, deflate_images=True, deflate_fonts=True)
    finally:
        pdf_doc.close()

    if is_same_file and os.path.exists(tmp_save_path):
        os.replace(tmp_save_path, out_file_path)

    yield json.dumps({
        "type": "done",
        "data": {
            "id": doc_id,
            "filename": os.path.basename(out_file_path),
            "filePath": out_file_path,
            "file_size": os.path.getsize(out_file_path),
            "num_pages": total_pages,
            "num_chunks": int(doc_record.get("num_chunks", 1)),
            "extracted_markdown": str(doc_record.get("extracted_markdown", "")),
            "status": "translated",
            "ingested_at": str(doc_record.get("ingested_at", "")),
            "fileType": "pdf"
        }
    }) + "\n"


def translate_document_inplace(
    doc_id: str,
    source_lang: str,
    target_lang: str,
    model: str = "llama3.2",
    backup_original: bool = True,
    target_dir: Optional[str] = None
) -> IngestResponse:
    """
    Dispatches translation to the DOCX or PDF fine-mode pipeline based on the document's
    stored file_type. Raises UnsupportedDocumentTypeError (mapped to HTTP 400) for any other
    type, ValueError (mapped to HTTP 404) if the document itself can't be found.
    """
    record = _load_doc_record(doc_id)
    file_type = record.get("file_type", "")
    if file_type == "docx":
        return translate_docx_inplace(doc_id, source_lang, target_lang, model, backup_original, target_dir)
    if file_type == "pdf":
        return translate_pdf_inplace_fine(doc_id, source_lang, target_lang, model, backup_original, target_dir)
    raise UnsupportedDocumentTypeError(
        f"In-place translation is not supported for file type '{file_type}'. Supported: docx, pdf."
    )

