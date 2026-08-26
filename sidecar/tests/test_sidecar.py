import os
import sys
import asyncio
import json
import pytest
from fastapi.testclient import TestClient
from starlette.requests import Request

# Ensure root workspace directory is in sys.path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))

from sidecar.main import app, global_exception_handler
from sidecar.domain.sanitizer import sanitize_extracted_text
from sidecar.domain.router import classify_file_type, DocumentCategory

client = TestClient(app)

def test_global_error_contract_is_safe_and_correlatable(monkeypatch):
    request = Request({
        "type": "http",
        "method": "POST",
        "path": "/contract-test",
        "query_string": b"",
        "headers": [],
        "scheme": "http",
        "server": ("testserver", 80),
        "client": ("testclient", 1),
    })
    logged = []
    monkeypatch.setattr("sidecar.main.logger.error", lambda message: logged.append(message))

    response = asyncio.run(global_exception_handler(request, RuntimeError("secret path C:/private/token")))
    payload = json.loads(response.body)
    event = json.loads(logged[0])

    assert response.status_code == 500
    assert payload["detail"] == "Internal Server Error"
    assert payload["error_id"] == event["error_id"]
    assert event == {
        "error_id": payload["error_id"],
        "error_type": "RuntimeError",
        "event": "unhandled_exception",
        "method": "POST",
        "path": "/contract-test",
    }
    assert "secret path" not in logged[0]

def test_health_endpoint():
    response = client.get("/health")
    assert response.status_code == 200
    data = response.json()
    assert data["status"] == "online"
    assert "engine" in data
    assert "version" in data
    assert "vector_db" in data
    assert "documents_count" in data
    assert "chunks_count" in data
    assert "ocr" in data
    assert "provider" in data["ocr"]
    assert "host_has_gpu" in data["ocr"]
    assert isinstance(data["ocr"]["host_has_gpu"], bool)


def test_text_sanitizer():
    dirty_text = "Hello\x00World\x07!\r\nLine 2\ufeff\n\n\n\n\nLine 3   "
    clean_text = sanitize_extracted_text(dirty_text)
    assert "\x00" not in clean_text
    assert "\x07" not in clean_text
    assert "\ufeff" not in clean_text
    assert "\r" not in clean_text
    assert "HelloWorld!" in clean_text
    assert "Line 2" in clean_text

def test_router_classification():
    assert classify_file_type("document.pdf") == DocumentCategory.PDF
    assert classify_file_type("photo.png") == DocumentCategory.IMAGE
    assert classify_file_type("report.docx") == DocumentCategory.DOCX
    assert classify_file_type("data.csv") == DocumentCategory.TABULAR
    assert classify_file_type("notes.md") == DocumentCategory.TEXT

def test_export_markdown_pdf():
    payload = {
        "markdown_content": "# Unit Test Document\n\nThis is a test document export.",
        "export_format": "pdf"
    }
    response = client.post("/export", json=payload)
    assert response.status_code == 200
    data = response.json()
    assert data["status"] == "success"
    assert "base64_content" in data
    assert "file_name" in data

def test_export_empty_markdown_raises_400():
    payload = {
        "markdown_content": "   ",
        "export_format": "pdf"
    }
    response = client.post("/export", json=payload)
    assert response.status_code == 400

def test_ingest_path_contract_rejects_invalid_limits_and_accepts_options(tmp_path):
    source = tmp_path / "document.pdf"
    source.write_bytes(b"not a real pdf")

    valid = client.post(
        "/ingest-path",
        json={
            "file_path": str(source),
            "vision_model": "llama3.2-vision",
            "vision_prompt": "Read the document",
            "normalize_with_llm": True,
            "normalization_model": "qwen2.5:7b",
            "max_tabular_rows": 300,
            "max_excel_rows_per_sheet": 150,
            "max_excel_sheets": 10,
        },
    )
    assert valid.status_code != 422

    invalid = client.post(
        "/ingest-path",
        json={"file_path": str(source), "max_excel_sheets": 0},
    )
    assert invalid.status_code == 422

def test_export_contract_rejects_blank_or_malformed_format():
    assert client.post("/export", json={"markdown_content": "   ", "export_format": "pdf"}).status_code == 400
    assert client.post("/export", json={"markdown_content": "# title", "export_format": "pdf/../../x"}).status_code == 422

def test_vector_search_endpoint():
    payload = {
        "query": "test query",
        "top_k": 3
    }
    response = client.post("/vector/search", json=payload)
    assert response.status_code == 200
    data = response.json()
    assert isinstance(data, list)

def test_vector_search_with_doc_ids():
    payload = {
        "query": "architecture overview",
        "top_k": 3,
        "doc_ids": ["doc-123", "doc-456"]
    }
    response = client.post("/vector/search", json=payload)
    assert response.status_code == 200
    data = response.json()
    assert isinstance(data, list)

def test_vector_search_empty_query():
    payload = {
        "query": "   ",
        "top_k": 3
    }
    response = client.post("/vector/search", json=payload)
    assert response.status_code == 200
    data = response.json()
    assert data == []

def test_vector_search_rejects_non_positive_top_k():
    response = client.post("/vector/search", json={"query": "test query", "top_k": 0})
    assert response.status_code == 422
    error = response.json()["detail"][0]
    assert error["loc"] == ["body", "top_k"]
    assert error["type"] == "greater_than_equal"

def test_tasks_cancel_endpoint():
    response = client.post("/tasks/cancel?task_id=test-123")
    assert response.status_code == 200
    data = response.json()
    assert data["status"] == "success"
    assert "test-123" in data["message"]

def test_cleanup_temp_endpoint():
    response = client.post("/cleanup/temp")
    assert response.status_code == 200
    data = response.json()
    assert data["status"] == "success"
    assert "cleaned_files" in data

def test_deterministic_fallback_embeddings():
    from sidecar.infrastructure.embeddings import get_fallback_embedding
    vec1 = get_fallback_embedding("Sample query for deterministic test", dim=384)
    vec2 = get_fallback_embedding("Sample query for deterministic test", dim=384)
    assert len(vec1) == 384
    assert len(vec2) == 384
    assert vec1 == vec2  # Exactly reproducible across runs

def test_history_index_search_and_project_filter():
    payload = {
        "id": "test-history-alpha",
        "session_id": "test-history-session-alpha",
        "project_path": "C:\\FakeTestProjects\\Alpha",
        "prompt": "Unit test prompt about database migration rollback strategy",
        "summary": "Testing history index",
        "outcome": "success",
        "started_at": "2026-01-01T00:00:00Z",
        "completed_at": "2026-01-01T00:05:00Z",
    }
    try:
        response = client.post("/history/index", json=payload)
        assert response.status_code == 200
        assert response.json()["success"] is True

        # Idempotent upsert: re-indexing the same id must not duplicate it in search results.
        assert client.post("/history/index", json=payload).status_code == 200

        search_response = client.post("/history/search", json={"query": "database migration rollback", "top_k": 20})
        assert search_response.status_code == 200
        results = search_response.json()
        matches = [r for r in results if r["id"] == "test-history-alpha"]
        assert len(matches) == 1
        assert matches[0]["project_path"] == payload["project_path"]
        assert matches[0]["score"] > 0

        scoped = client.post("/history/search", json={
            "query": "database migration rollback",
            "top_k": 20,
            "project_paths": ["C:\\FakeTestProjects\\Alpha"],
        })
        assert any(r["id"] == "test-history-alpha" for r in scoped.json())

        excluded = client.post("/history/search", json={
            "query": "database migration rollback",
            "top_k": 20,
            "project_paths": ["C:\\FakeTestProjects\\SomethingElse"],
        })
        assert all(r["id"] != "test-history-alpha" for r in excluded.json())
    finally:
        client.post("/history/remove", json={"session_ids": ["test-history-session-alpha"]})

def test_history_remove_by_project():
    payload = {
        "id": "test-history-beta",
        "session_id": "test-history-session-beta",
        "project_path": "C:\\FakeTestProjects\\Beta",
        "prompt": "Unit test prompt for project-scoped removal",
        "outcome": "success",
        "started_at": "2026-01-01T00:00:00Z",
    }
    client.post("/history/index", json=payload)
    remove_response = client.post("/history/remove", json={"project_path": "C:\\FakeTestProjects\\Beta"})
    assert remove_response.status_code == 200
    search_response = client.post("/history/search", json={"query": "project-scoped removal", "top_k": 20})
    assert all(r["id"] != "test-history-beta" for r in search_response.json())

def test_history_index_rejects_malformed_id():
    payload = {
        "id": "bad\"id",
        "session_id": "test-history-session-gamma",
        "project_path": "C:\\FakeTestProjects\\Gamma",
        "prompt": "test",
        "outcome": "success",
        "started_at": "2026-01-01T00:00:00Z",
    }
    response = client.post("/history/index", json=payload)
    assert response.status_code == 400

def test_history_search_empty_query_returns_empty_list():
    response = client.post("/history/search", json={"query": "   "})
    assert response.status_code == 200
    assert response.json() == []

def test_ocr_prepare_image_resizing():
    from sidecar.infrastructure.ocr import _prepare_image_for_ocr
    try:
        from PIL import Image
        import io
        
        # Create large 3000x2000 image
        large_img = Image.new("RGB", (3000, 2000), color="blue")
        buf = io.BytesIO()
        large_img.save(buf, format="PNG")
        raw_bytes = buf.getvalue()
        
        resized_bytes = _prepare_image_for_ocr(raw_bytes, max_dim=1500)
        resized_img = Image.open(io.BytesIO(resized_bytes))
        assert max(resized_img.size) <= 1500
    except ImportError:
        # Gracefully test passthrough fallback when PIL is not present
        dummy_bytes = b"non-image-binary-payload"
        assert _prepare_image_for_ocr(dummy_bytes, max_dim=1500) == dummy_bytes

def test_rapid_ocr_extracts_text_from_image():
    """RapidOCR (Tier 1, fast local GPU-capable engine) must actually recognize rendered text."""
    from sidecar.infrastructure.ocr import run_rapid_ocr
    from PIL import Image, ImageDraw
    import io

    img = Image.new("RGB", (400, 100), color="white")
    draw = ImageDraw.Draw(img)
    draw.text((10, 10), "ONLYRAG INVOICE TOTAL 4200", fill="black")
    buf = io.BytesIO()
    img.save(buf, format="PNG")

    text_out = run_rapid_ocr(buf.getvalue())
    assert "ONLY" in text_out.upper() and "RAG" in text_out.upper()
    assert "INVOICE" in text_out.upper()
    assert "4200" in text_out

def test_run_layout_ocr_uses_rapidocr(monkeypatch):
    """run_layout_ocr must directly delegate to RapidOCR on rendered images."""
    from sidecar.infrastructure import ocr as ocr_module
    called = []
    def fake_rapid(img_bytes):
        called.append(img_bytes)
        return "Fast tier only please"
    monkeypatch.setattr(ocr_module, "run_rapid_ocr", fake_rapid)
    result = ocr_module.run_layout_ocr(b"dummy_test_bytes")
    assert "Fast tier only please".lower().split()[0] in result.lower()
    assert len(called) == 1

def test_ocr_vision_fallback_sets_keep_alive_zero():
    """run_vision_ocr must evict the vision model immediately after use (keep_alive: 0),
    per the documented Ephemeral Eviction policy for OCR support models, to avoid VRAM
    thrashing against the pinned primary model."""
    from sidecar.infrastructure import ocr as ocr_module

    captured_payloads = []

    class FakeResponse:
        status_code = 200
        def json(self):
            return {"response": "Extracted text"}

    class FakeHttpxClient:
        def post(self, url, json=None, timeout=None):
            captured_payloads.append(json)
            return FakeResponse()

    original_client = ocr_module.httpx_client
    ocr_module.httpx_client = FakeHttpxClient()
    try:
        result = ocr_module.run_vision_ocr(b"fake-image-bytes", model="llama3.2-vision")
    finally:
        ocr_module.httpx_client = original_client

    assert result == "Extracted text"
    assert captured_payloads, "Expected at least one call to the Ollama /api/generate endpoint"
    assert captured_payloads[0]["keep_alive"] == 0

def test_extract_pdf_document_native_text_page():
    """extract_pdf_document must extract native page text and report the correct page number."""
    import pymupdf
    from sidecar.domain.ingestion import extract_pdf_document

    doc = pymupdf.open()
    try:
        page = doc.new_page(width=595, height=842)
        page.insert_text((50, 72), "Hello OnlyRag native text extraction page content", fontsize=12)

        pages = extract_pdf_document(doc)
        assert len(pages) == 1
        page_num, page_text = pages[0]
        assert page_num == 1
        assert "Hello OnlyRag native text extraction page content" in page_text
    finally:
        doc.close()

def test_extract_pdf_document_parallelizes_ocr_pages(monkeypatch):
    """extract_pdf_document must run the OCR rendering phase concurrently (bounded by
    PDF_PAGE_RENDER_CONCURRENCY) instead of one page at a time, and still return pages in the
    correct order despite the underlying work overlapping."""
    import time
    import pymupdf
    from sidecar.domain import ingestion as ingestion_module

    def slow_ocr(image_bytes):
        time.sleep(0.3)
        return "OCR page text"
    monkeypatch.setattr(ingestion_module, "run_layout_ocr", slow_ocr)

    doc = pymupdf.open()
    try:
        for _ in range(6):
            doc.new_page(width=200, height=200)  # blank page -> under the 40-char threshold -> OCR_REQUIRED

        start = time.monotonic()
        pages = ingestion_module.extract_pdf_document(doc)
        elapsed = time.monotonic() - start

        assert [p[0] for p in pages] == [1, 2, 3, 4, 5, 6], "Page order must be preserved despite concurrent rendering"
        assert all("OCR page text" in p[1] for p in pages)
        assert elapsed < 1.2, f"Expected concurrent OCR rendering to overlap, but took {elapsed:.2f}s (sequential would be ~1.8s)"
    finally:
        doc.close()

def _render_pdf_page_content(doc, page, page_num, raw_text, md_tables, used_ocr, **kwargs):
    from sidecar.domain.ingestion import prepare_pdf_page_work_item, render_prepared_pdf_page
    work_item = prepare_pdf_page_work_item(
        doc, page, page_num, raw_text, md_tables, used_ocr,
        vision_model=kwargs.get("vision_model"),
        vision_prompt=kwargs.get("vision_prompt"),
        filename=kwargs.get("filename", ""),
        num_pages=kwargs.get("num_pages", 1)
    )
    _, content = render_prepared_pdf_page(work_item)
    return content


def test_render_pdf_page_content_ocr_path(monkeypatch):
    """prepare_pdf_page_work_item + render_prepared_pdf_page must route to run_layout_ocr when used_ocr=True."""
    import pymupdf
    from sidecar.domain import ingestion as ingestion_module

    monkeypatch.setattr(ingestion_module, "run_layout_ocr", lambda img_bytes: "OCR-extracted markdown")

    doc = pymupdf.open()
    try:
        page = doc.new_page(width=595, height=842)
        result = _render_pdf_page_content(
            doc, page, page_num=1, raw_text="", md_tables=[], used_ocr=True
        )
        assert "OCR-extracted markdown" in result
    finally:
        doc.close()

def _new_scanned_pdf_page(doc):
    """Test helper: a scanned PDF page with 0 native text and a rendered bitmap text image."""
    from PIL import Image, ImageDraw
    import pymupdf
    import io

    page = doc.new_page(width=595, height=842)
    img = Image.new("RGB", (800, 300), color="white")
    draw = ImageDraw.Draw(img)
    draw.text((20, 50), "SCANNED INVOICE ITEM TOTAL 9900 EUR", fill="black")
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    page.insert_image(page.rect, stream=buf.getvalue())
    return page

def test_scanned_pdf_page_ocr_extraction(monkeypatch):
    """A scanned PDF with an embedded bitmap text layer must have its text detected by OCR."""
    import pymupdf
    from sidecar.domain import ingestion as ingestion_module

    monkeypatch.setattr(ingestion_module, "run_layout_ocr", lambda img: "SCANNED INVOICE ITEM TOTAL 9900 EUR")

    doc = pymupdf.open()
    try:
        _new_scanned_pdf_page(doc)
        pages = ingestion_module.extract_pdf_document(doc)
        assert len(pages) == 1
        page_num, page_text = pages[0]
        assert page_num == 1
        assert "SCANNED INVOICE" in page_text.upper() or "9900" in page_text
    finally:
        doc.close()

def test_analyze_pdf_page_structure_classifies_scanned_page_as_ocr_required():
    """A page with image content and zero text must be classified OCR_REQUIRED."""
    import pymupdf
    from sidecar.domain.router import analyze_pdf_page_structure, PageRoutingStrategy

    doc = pymupdf.open()
    try:
        page = _new_scanned_pdf_page(doc)
        struct_info = analyze_pdf_page_structure(page)
        assert struct_info["strategy"] == PageRoutingStrategy.OCR_REQUIRED
    finally:
        doc.close()

def test_analyze_pdf_page_structure_classifies_digital_page_as_native_text():
    """A page with substantial native text must be classified NATIVE_TEXT."""
    import pymupdf
    from sidecar.domain.router import analyze_pdf_page_structure, PageRoutingStrategy

    doc = pymupdf.open()
    try:
        page = doc.new_page(width=595, height=842)
        page.insert_text((50, 72), "Digital PDF page content with sufficient character length for native text.", fontsize=12)
        struct_info = analyze_pdf_page_structure(page)
        assert struct_info["strategy"] == PageRoutingStrategy.NATIVE_TEXT
    finally:
        doc.close()

def test_analyze_pdf_page_structure_keeps_short_image_free_text_native():
    """A page with a short image-free native text layer must be classified NATIVE_TEXT."""
    import pymupdf
    from sidecar.domain.router import analyze_pdf_page_structure, PageRoutingStrategy

    doc = pymupdf.open()
    try:
        page = doc.new_page(width=595, height=842)
        page.insert_text((50, 72), "日本語", fontsize=12)
        struct_info = analyze_pdf_page_structure(page)
        assert struct_info["strategy"] == PageRoutingStrategy.NATIVE_TEXT
    finally:
        doc.close()

def test_analyze_pdf_page_structure_blank_page_stays_ocr_required():
    """A blank page must still route to OCR_REQUIRED."""
    import pymupdf
    from sidecar.domain.router import analyze_pdf_page_structure, PageRoutingStrategy

    doc = pymupdf.open()
    try:
        page = doc.new_page(width=200, height=200)
        struct_info = analyze_pdf_page_structure(page)
        assert struct_info["strategy"] == PageRoutingStrategy.OCR_REQUIRED
    finally:
        doc.close()

def test_extract_tabular_document_csv_and_json():
    from sidecar.domain.ingestion import extract_tabular_document
    csv_bytes = b"ColA,ColB,ColC\nVal1,Val2,Val3\nVal4,Val5|Pipe,Val6"
    csv_res = extract_tabular_document("data.csv", csv_bytes, None)
    assert len(csv_res) == 1
    assert "ColA" in csv_res[0][1] and "ColB" in csv_res[0][1] and "ColC" in csv_res[0][1]
    assert "Val5" in csv_res[0][1] and "Pipe" in csv_res[0][1]

    json_bytes = b'{"name": "OnlyRag", "version": 2}'
    json_res = extract_tabular_document("config.json", json_bytes, None)
    assert len(json_res) == 1
    assert "```json" in json_res[0][1]

def test_extract_tabular_document_truncation_note():
    from sidecar.domain.ingestion import extract_tabular_document
    # CSV with 10 rows, limit to 4
    lines = ["ColA,ColB"] + [f"Row{i},Val{i}" for i in range(10)]
    csv_bytes = "\n".join(lines).encode("utf-8")
    res = extract_tabular_document("big.csv", csv_bytes, None, max_rows=4)
    assert len(res) == 1
    text = res[0][1]
    assert "Tabella troncata a 4 righe (su 10 totali" in text
    assert "Row0" in text and "Val0" in text
    assert "Row3" in text and "Val3" in text
    assert "Row4" not in text

def test_generate_embedding_with_status_tracks_fallback():
    from sidecar.infrastructure.embeddings import generate_embedding_with_status
    vec, is_fallback = generate_embedding_with_status("Test sentence for embedding vector")
    assert isinstance(vec, list)
    assert len(vec) in (384, 768)
    assert isinstance(is_fallback, bool)

def test_docx_image_extraction_ocr(tmp_path, monkeypatch):
    import docx
    from PIL import Image, ImageDraw
    import io
    from sidecar.domain import ingestion as ingestion_module

    monkeypatch.setattr(ingestion_module, "run_layout_ocr", lambda img: "DOCX IMAGE TEXT 7788")

    # Create image with text
    img = Image.new("RGB", (300, 80), color="white")
    draw = ImageDraw.Draw(img)
    draw.text((10, 10), "DOCX IMAGE TEXT 7788", fill="black")
    img_buf = io.BytesIO()
    img.save(img_buf, format="PNG")
    img_bytes = img_buf.getvalue()

    # Create docx with paragraph and image
    doc = docx.Document()
    doc.add_paragraph("Paragraph inside docx document")
    img_path = str(tmp_path / "temp_img.png")
    with open(img_path, "wb") as f:
        f.write(img_bytes)
    doc.add_picture(img_path)
    
    docx_path = str(tmp_path / "doc_with_image.docx")
    doc.save(docx_path)

    with open(docx_path, "rb") as f:
        docx_bytes = f.read()

    md, num_pages = ingestion_module.extract_document_markdown("doc_with_image.docx", docx_bytes, docx_path)
    assert "Paragraph inside docx document" in md
    assert "DOCX IMAGE TEXT" in md or "7788" in md

def test_semantic_chunks_oversized_line_splitting():
    from sidecar.domain.ingestion import create_semantic_chunks
    long_line = "This is a very long text sentence. " * 100  # ~3500 chars
    md = f"# LargeDoc\n\n## Section 1\n\n{long_line}"
    chunks = create_semantic_chunks("LargeDoc.md", md)
    assert len(chunks) > 1
    for idx, c_text, header in chunks:
        assert len(c_text) < 1500
        assert "LargeDoc" in header
        assert "[Documento: LargeDoc.md | Sezione:" in c_text


def test_reciprocal_rank_fusion_k60():
    from sidecar.services.search_service import reciprocal_rank_fusion
    dense_ranks = {"chunk_1": 1, "chunk_2": 2, "chunk_3": 3}
    sparse_ranks = {"chunk_2": 1, "chunk_1": 2, "chunk_4": 3}

    rrf = reciprocal_rank_fusion(dense_ranks, sparse_ranks, k=60)

    # chunk_1: 1/(60+1) + 1/(60+2) = 1/61 + 1/62 = 0.016393 + 0.016129 = 0.032522
    # chunk_2: 1/(60+2) + 1/(60+1) = 0.032522
    # chunk_3: 1/(60+3) = 1/63 = 0.015873
    # chunk_4: 1/(60+3) = 1/63 = 0.015873
    assert pytest.approx(rrf["chunk_1"], rel=1e-3) == (1/61 + 1/62)
    assert pytest.approx(rrf["chunk_2"], rel=1e-3) == (1/61 + 1/62)
    assert pytest.approx(rrf["chunk_3"], rel=1e-3) == (1/63)
    assert pytest.approx(rrf["chunk_4"], rel=1e-3) == (1/63)
    assert rrf["chunk_1"] > rrf["chunk_3"]


def test_reranker_cross_scoring():
    from sidecar.infrastructure.reranker import calculate_cross_score, rerank_candidates
    query = "quantum computing error correction"
    relevant_text = "This paper investigates quantum computing error correction algorithms in superconducting qubits."
    irrelevant_text = "The recipe for pasta includes flour, eggs, and a pinch of salt."

    score_rel = calculate_cross_score(query, relevant_text, header="Quantum Physics")
    score_irrel = calculate_cross_score(query, irrelevant_text, header="Cooking")

    assert score_rel > score_irrel
    assert score_rel > 0.5
    assert score_irrel < 0.2

    candidates = [
        {"chunk_id": "c_irr", "text": irrelevant_text, "score": 0.6, "doc_name": "cook.md"},
        {"chunk_id": "c_rel", "text": relevant_text, "score": 0.4, "doc_name": "quantum.md"}
    ]

    reranked = rerank_candidates(query, candidates, top_k=2)
    assert len(reranked) == 2
    assert reranked[0]["chunk_id"] == "c_rel"
    assert reranked[0]["score"] > reranked[1]["score"]


def test_word_segmenter_defragment_and_compounds():
    from sidecar.domain.word_segmenter import normalize_ocr_token_spacing

    # Test spaced letter de-fragmentation
    spaced_raw = "LOCALITA * BORGO MA N TO VA NO Lo C.Villa Poma"
    normalized_spaced = normalize_ocr_token_spacing(spaced_raw)
    assert "MANTOVANO" in normalized_spaced

    # Test compound with apostrophe segmentation
    compound_raw = "Aseguitodell'eserciziodeldirittodirecessodal Contratto Telepass Family"
    normalized_compound = normalize_ocr_token_spacing(compound_raw)
    assert "seguito" in normalized_compound.lower()
    assert "esercizio" in normalized_compound.lower()
    assert "recesso" in normalized_compound.lower()

    # Test preposition and phrase splitting
    fused_raw = "conlapresente chiedelacessazione delcontratto sopraindicato viadel Serafico"
    normalized_fused = normalize_ocr_token_spacing(fused_raw)
    assert "con la presente" in normalized_fused
    assert "chiede la cessazione" in normalized_fused
    assert "del contratto" in normalized_fused
    assert "sopra indicato" in normalized_fused
    assert "via del" in normalized_fused


def test_multilang_vocab_manager_and_segmentation():
    from sidecar.domain.word_segmenter import normalize_language_code, get_vocab_manager, normalize_ocr_token_spacing

    assert normalize_language_code("Italian") == "it"
    assert normalize_language_code("eng") == "en"
    assert normalize_language_code("Spanish") == "es"
    assert normalize_language_code("fr_FR") == "fr"
    assert normalize_language_code(None) == "it"

    mgr = get_vocab_manager()
    assert mgr.is_known_word("contratto", "it")
    assert mgr.is_known_word("contract", "en")
    assert mgr.get_word_zipf("casa", "it") > 0.0

    # Test multi-language token spacing with language parameter
    en_fused = "consequentderegistration of the equipment"
    en_norm = normalize_ocr_token_spacing(en_fused, lang="en")
    assert "equipment" in en_norm.lower()


def test_vocab_sync_service_offline_and_caching(tmp_path):
    import json
    import pytest
    from sidecar.services.vocab_service import VocabSyncService
    from sidecar.domain.word_segmenter import MultiLangVocabManager

    cache_dir = str(tmp_path / "vocab_cache")
    svc = VocabSyncService(manifest_url="http://127.0.0.1:9999/nonexistent/manifest.json", cache_dir=cache_dir)
    
    # Run sync against offline URL
    import asyncio
    res = asyncio.run(svc.sync_vocabularies(timeout_sec=0.5))
    assert res["status"] in ("offline", "cached")

    # Verify custom local dictionary loading in MultiLangVocabManager
    custom_vocab = {"personalizzato": 6.8, "ultratecnico": 7.2}
    with open(tmp_path / "vocab_cache" / "it.json", "w", encoding="utf-8") as f:
        json.dump(custom_vocab, f)

    mgr = MultiLangVocabManager(cache_dir=cache_dir)
    assert mgr.get_word_zipf("personalizzato", "it") == pytest.approx(6.8)
    assert mgr.get_word_zipf("ultratecnico", "it") == pytest.approx(7.2)


def test_vocab_status_and_sync_endpoints():
    status_resp = client.get("/vocab/status")
    assert status_resp.status_code == 200
    data = status_resp.json()
    assert "wordfreq_available" in data
    assert "cached_languages" in data

    sync_resp = client.post("/vocab/sync")
    assert sync_resp.status_code == 200
    sync_data = sync_resp.json()
    assert "status" in sync_data


def test_opencv_deskew():
    import numpy as np
    import cv2
    from sidecar.infrastructure.ocr import compute_deskew_angle, deskew_image

    # Create a synthetic white image with black text rectangle
    img = np.ones((200, 400, 3), dtype=np.uint8) * 255
    cv2.rectangle(img, (50, 80), (350, 120), (0, 0, 0), -1)
    is_success, buf = cv2.imencode(".png", img)
    assert is_success
    png_bytes = buf.tobytes()

    # Test angle calculation on horizontal image
    angle = compute_deskew_angle(img)
    assert abs(angle) < 1.0

    # Test deskew function returns valid PNG bytes
    deskewed_bytes = deskew_image(png_bytes)
    assert len(deskewed_bytes) > 0


def test_language_detection_and_target_skip():
    from sidecar.domain.translator import detect_block_language, is_block_in_target_lang

    it_text = "Questo documento contiene i dati relativi alla sezione contrattuale."
    en_text = "This document contains important data regarding the contract section."
    zh_text = "这是一个包含合同条款的重要文件。"

    assert detect_block_language(it_text) == "italian"
    assert detect_block_language(en_text) == "english"
    assert detect_block_language(zh_text) == "chinese"

    assert is_block_in_target_lang(it_text, "Italian") is True
    assert is_block_in_target_lang(it_text, "English") is False
    assert is_block_in_target_lang(en_text, "English") is True
    assert is_block_in_target_lang(zh_text, "Chinese") is True


def test_output_path_resolution(tmp_path):
    from sidecar.domain.translator import _resolve_output_filepath

    src_file = str(tmp_path / "original_document.pdf")
    with open(src_file, "w") as f:
        f.write("dummy")

    target_dir = str(tmp_path / "custom_exports")
    resolved = _resolve_output_filepath(src_file, "original_document.pdf", "Italian", target_dir)
    assert resolved.endswith("original_document_italian.pdf")
    assert "custom_exports" in resolved
    assert os.path.exists(target_dir)


def test_translate_inplace_stream_endpoint_404():
    payload = {
        "source_lang": "Italian",
        "target_lang": "English"
    }
    response = client.post("/documents/non-existent-doc-9999/translate-inplace-stream", json=payload)
    assert response.status_code == 200
    assert "error" in response.text.lower() or "not found" in response.text.lower()



if __name__ == "__main__":
    import inspect
    import sys
    current_module = sys.modules[__name__]
    test_functions = [obj for name, obj in inspect.getmembers(current_module) if inspect.isfunction(obj) and name.startswith("test_")]
    passed = 0
    failed = 0
    print(f"Running {len(test_functions)} sidecar test functions...")
    for fn in test_functions:
        try:
            fn()
            passed += 1
            print(f"  [PASS] {fn.__name__}")
        except Exception as e:
            failed += 1
            print(f"  [FAIL] {fn.__name__}: {e}")
    print(f"\nResult: {passed} passed, {failed} failed.")
    if failed > 0:
        sys.exit(1)


def test_list_stored_documents_includes_fallback_embedded_documents(monkeypatch):
    """Documents indexed with fallback embeddings must stay visible and be flagged as such."""
    import sidecar.services.search_service as search_service

    stored_rows = [
        {"id": "doc-1", "filename": "regular.pdf", "status": "indexed", "num_chunks": 4},
        {"id": "doc-2", "filename": "degraded.pdf", "status": "indexed_fallback", "num_chunks": 2},
        {"id": "doc-3", "filename": "broken.pdf", "status": "failed", "num_chunks": 0},
    ]

    class FakeTable:
        def to_arrow(self):
            raise RuntimeError("arrow unavailable")

        def to_pandas(self):
            class FakeFrame:
                def to_dict(self, orient="records"):
                    return stored_rows
            return FakeFrame()

    class FakeDb:
        def open_table(self, name):
            return FakeTable()

    monkeypatch.setattr(search_service, "get_existing_tables", lambda: [search_service.DOCS_TABLE_NAME])
    monkeypatch.setattr(search_service, "lance_db", FakeDb())

    listed = search_service.list_stored_documents()
    by_id = {d["id"]: d for d in listed}

    assert set(by_id) == {"doc-1", "doc-2"}, "failed documents must stay hidden, fallback ones must not"
    assert by_id["doc-1"]["status"] == "indexed"
    assert by_id["doc-1"]["used_fallback_embeddings"] is False
    assert by_id["doc-2"]["status"] == "indexed_fallback"
    assert by_id["doc-2"]["used_fallback_embeddings"] is True

# ---------------------------------------------------------------------------
# Vision LLM OCR engine routing
# ---------------------------------------------------------------------------

def test_render_vision_prompt_fills_per_page_variables():
    """The per-page Mustache variables of `images:analysis` are only known inside the sidecar page
    loop, so the renderer ships the raw template and this is where it gets filled."""
    from sidecar.domain.vision_prompt import render_vision_prompt

    template = "Doc {{filename}} page {{currentPage}}/{{numPages}}\nContext: {{activePageContent}}"
    out = render_vision_prompt(template, "report.pdf", 3, 8, "Operating margin +24%")

    assert "Doc report.pdf page 3/8" in out
    assert "Operating margin +24%" in out
    assert "{{" not in out

def test_render_vision_prompt_falls_back_when_template_empty():
    """An override the user emptied must not send a blank instruction to the vision model."""
    from sidecar.domain.vision_prompt import render_vision_prompt, FALLBACK_VISION_PROMPT

    assert render_vision_prompt("   ", "a.pdf", 1, 1) == FALLBACK_VISION_PROMPT
    assert render_vision_prompt(None, "a.pdf", 1, 1) == FALLBACK_VISION_PROMPT

def test_run_page_ocr_stays_on_rapidocr_without_vision_prompt(monkeypatch):
    """No vision prompt on the wire means the user kept the native CUDA engine."""
    from sidecar.domain import ingestion as ingestion_module

    monkeypatch.setattr(ingestion_module, "run_layout_ocr", lambda img: "RAPIDOCR TEXT")
    monkeypatch.setattr(ingestion_module, "run_vision_ocr", lambda *a, **kw: pytest.fail(
        "Vision model must not be called when no vision prompt was requested"))

    assert ingestion_module.run_page_ocr(b"img", vision_model="llama3.2-vision") == "RAPIDOCR TEXT"

def test_run_page_ocr_routes_to_vision_with_rendered_prompt(monkeypatch):
    """With a vision prompt present the page bitmap goes to the multimodal model, and the prompt
    reaching it is the rendered one, not the raw template."""
    from sidecar.domain import ingestion as ingestion_module

    captured = {}
    def fake_vision(image_bytes, prompt, ollama_url=None, model=None):
        captured["prompt"] = prompt
        captured["model"] = model
        return "VISION MARKDOWN"

    monkeypatch.setattr(ingestion_module, "run_vision_ocr", fake_vision)
    monkeypatch.setattr(ingestion_module, "run_layout_ocr", lambda img: pytest.fail(
        "RapidOCR must not run when the vision model returned content"))

    result = ingestion_module.run_page_ocr(
        b"img",
        vision_model="minicpm-v",
        vision_prompt="Read page {{currentPage}} of {{numPages}} in {{filename}}",
        filename="invoice.pdf",
        page_num=2,
        num_pages=5
    )

    assert result == "VISION MARKDOWN"
    assert captured["prompt"] == "Read page 2 of 5 in invoice.pdf"
    assert captured["model"] == "minicpm-v"

def test_run_page_ocr_falls_back_to_rapidocr_when_vision_returns_nothing(monkeypatch):
    """run_vision_ocr returns an empty string on a missing model, a timeout or a refusal — a page
    must never be silently dropped over that."""
    from sidecar.domain import ingestion as ingestion_module

    monkeypatch.setattr(ingestion_module, "run_vision_ocr", lambda *a, **kw: "")
    monkeypatch.setattr(ingestion_module, "run_layout_ocr", lambda img: "RAPIDOCR SAFETY NET")

    result = ingestion_module.run_page_ocr(
        b"img", vision_model="llama3.2-vision", vision_prompt="Transcribe this page"
    )
    assert result == "RAPIDOCR SAFETY NET"

def test_render_pdf_page_content_routes_scanned_page_to_vision(monkeypatch):
    """End of the wire: a scanned PDF page prepared with a vision prompt is transcribed by the
    multimodal model, with the document filename and page count reaching the prompt."""
    from sidecar.domain import ingestion as ingestion_module

    captured = {}
    def fake_vision(image_bytes, prompt, ollama_url=None, model=None):
        captured["prompt"] = prompt
        return "# Scanned Page\n\nVision transcription"

    monkeypatch.setattr(ingestion_module, "run_vision_ocr", fake_vision)

    import pymupdf

    doc = pymupdf.open()
    doc.new_page(width=595, height=842)
    try:
        page = doc.load_page(0)
        content = _render_pdf_page_content(
            doc, page, 1, "", [], True,
            vision_model="llama3.2-vision",
            vision_prompt="File {{filename}}, page {{currentPage}}/{{numPages}}",
            filename="scan.pdf",
            num_pages=4
        )
    finally:
        doc.close()

    assert "Vision transcription" in content
    assert captured["prompt"] == "File scan.pdf, page 1/4"

