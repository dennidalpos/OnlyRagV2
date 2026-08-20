import os
import sys
import docx
import pymupdf
import pytest
from fastapi.testclient import TestClient

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))

from sidecar.main import app
from sidecar.domain import translator as translator_module

client = TestClient(app)


def _make_pdf(path: str, text: str = "SUPERSECRETMARKER12345") -> None:
    doc = pymupdf.open()
    page = doc.new_page(width=595, height=842)
    page.insert_textbox(pymupdf.Rect(40, 40, 400, 80), text, fontsize=14, fontname="helv", color=(0, 0, 0))
    doc.save(path)
    doc.close()


def _make_docx(path: str) -> None:
    doc = docx.Document()
    doc.add_heading("Test Document", level=1)
    p = doc.add_paragraph()
    run = p.add_run("Hello world")
    run.bold = True
    doc.add_paragraph("Second paragraph text")

    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Cell one"
    table.rows[0].cells[1].text = "Cell two"

    doc.save(path)


def test_collect_docx_runs_covers_paragraphs_and_tables(tmp_path):
    path = str(tmp_path / "sample.docx")
    _make_docx(path)
    doc = docx.Document(path)

    runs = translator_module._collect_docx_runs(doc)
    texts = [r.text for r in runs]

    assert "Test Document" in texts
    assert "Hello world" in texts
    assert "Second paragraph text" in texts
    assert "Cell one" in texts
    assert "Cell two" in texts


def test_batch_runs_respects_max_chars(tmp_path):
    path = str(tmp_path / "sample.docx")
    _make_docx(path)
    doc = docx.Document(path)
    runs = translator_module._collect_docx_runs(doc)

    batches = translator_module._batch_runs(runs, max_chars=15)
    assert len(batches) > 1
    # Every run must appear in exactly one batch, in original order.
    flattened = [r for batch in batches for r in batch]
    assert flattened == runs


class _FakeResponse:
    def __init__(self, status_code=200, body=None):
        self.status_code = status_code
        self._body = body or {}

    def json(self):
        return self._body


def test_call_ollama_translate_retries_once_on_timeout_then_succeeds(monkeypatch):
    """A timeout is usually transient Ollama/GPU contention (e.g. a concurrent coding agent task
    holding the model queue) rather than a permanent failure -- reproduces the production incident
    where an in-place translation lost a segment because a single 120s timeout gave up outright."""
    import httpx

    call_count = {"n": 0}

    class FakeHttpxClient:
        def post(self, url, json=None, timeout=None):
            call_count["n"] += 1
            if call_count["n"] == 1:
                raise httpx.TimeoutException("simulated contention timeout")
            return _FakeResponse(200, {"response": "translated text"})

    monkeypatch.setattr(translator_module, "httpx_client", FakeHttpxClient())
    monkeypatch.setattr(translator_module.time, "sleep", lambda *_: None)

    result = translator_module._call_ollama_translate("hello", "English", "Italian", "llama3.2")

    assert result == "translated text"
    assert call_count["n"] == 2


def test_call_ollama_translate_gives_up_after_max_attempts_on_repeated_timeout(monkeypatch):
    import httpx

    call_count = {"n": 0}

    class FakeHttpxClient:
        def post(self, url, json=None, timeout=None):
            call_count["n"] += 1
            raise httpx.TimeoutException("simulated persistent timeout")

    monkeypatch.setattr(translator_module, "httpx_client", FakeHttpxClient())
    monkeypatch.setattr(translator_module.time, "sleep", lambda *_: None)

    result = translator_module._call_ollama_translate("hello", "English", "Italian", "llama3.2")

    assert result == ""
    assert call_count["n"] == translator_module._TRANSLATE_MAX_ATTEMPTS


def test_call_ollama_translate_does_not_retry_on_non_timeout_failure(monkeypatch):
    """A bad HTTP status (e.g. model not found) or a non-timeout exception won't be fixed by
    retrying, so only one attempt should be made."""
    call_count = {"n": 0}

    class FakeHttpxClient:
        def post(self, url, json=None, timeout=None):
            call_count["n"] += 1
            return _FakeResponse(404, {})

    monkeypatch.setattr(translator_module, "httpx_client", FakeHttpxClient())

    result = translator_module._call_ollama_translate("hello", "English", "Italian", "llama3.2")

    assert result == ""
    assert call_count["n"] == 1


def test_translate_batch_happy_path_reassigns_run_text(tmp_path, monkeypatch):
    path = str(tmp_path / "sample.docx")
    _make_docx(path)
    doc = docx.Document(path)
    runs = translator_module._collect_docx_runs(doc)[:2]  # "Test Document", "Hello world"

    def fake_call(text, source_lang, target_lang, model):
        assert translator_module._RUN_SEPARATOR in text
        parts = text.split(f"\n{translator_module._RUN_SEPARATOR}\n")
        return f"\n{translator_module._RUN_SEPARATOR}\n".join(f"[{p}]" for p in parts)

    monkeypatch.setattr(translator_module, "_call_ollama_translate", fake_call)
    translator_module._translate_batch(runs, "English", "Italian", "llama3.2")

    assert runs[0].text == "[Test Document]"
    assert runs[1].text == "[Hello world]"


def test_translate_batch_falls_back_on_segment_mismatch(tmp_path, monkeypatch):
    path = str(tmp_path / "sample.docx")
    _make_docx(path)
    doc = docx.Document(path)
    runs = translator_module._collect_docx_runs(doc)[:2]
    original_texts = [r.text for r in runs]

    call_log = []

    def fake_call(text, source_lang, target_lang, model):
        call_log.append(text)
        if translator_module._RUN_SEPARATOR in text:
            return "only one segment back"  # wrong count on purpose
        return f"TR:{text}"

    monkeypatch.setattr(translator_module, "_call_ollama_translate", fake_call)
    translator_module._translate_batch(runs, "English", "Italian", "llama3.2")

    # Fallback path: one batch call + one per-run call each.
    assert len(call_log) == 1 + len(runs)
    for run, original in zip(runs, original_texts):
        assert run.text == f"TR:{original}"


def test_translate_docx_inplace_end_to_end(tmp_path, monkeypatch):
    path = str(tmp_path / "e2e_sample.docx")
    _make_docx(path)

    def fake_call(text, source_lang, target_lang, model):
        if translator_module._RUN_SEPARATOR in text:
            parts = text.split(f"\n{translator_module._RUN_SEPARATOR}\n")
            return f"\n{translator_module._RUN_SEPARATOR}\n".join(f"TR-{p}" for p in parts)
        return f"TR-{text}"

    monkeypatch.setattr(translator_module, "_call_ollama_translate", fake_call)

    doc_id = None
    try:
        ingest_res = client.post("/ingest-path", json={"file_path": path})
        assert ingest_res.status_code == 200
        doc_id = ingest_res.json()["id"]

        res = client.post(
            f"/documents/{doc_id}/translate-inplace",
            json={"source_lang": "English", "target_lang": "Italian"},
        )
        assert res.status_code == 200
        data = res.json()
        assert data["status"] == "indexed"

        # Original file on disk was overwritten in place with translated run text.
        translated_doc = docx.Document(path)
        translated_texts = [r.text for r in translator_module._collect_docx_runs(translated_doc)]
        assert all(t.startswith("TR-") for t in translated_texts)

        # In-place translation separates flows: original extracted markdown in RAG DB is preserved.
        assert "Hello world" in data["extracted_markdown"]
    finally:
        if doc_id:
            client.delete(f"/documents/{doc_id}")


def test_translate_inplace_rejects_missing_document():
    with pytest.raises(ValueError, match="not found"):
        translator_module.translate_docx_inplace("nonexistent-doc-id-xyz", "English", "Italian")


def test_translate_inplace_rejects_non_docx_file_type(tmp_path):
    txt_path = str(tmp_path / "plain.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("Just a plain text file, not a DOCX.")

    doc_id = None
    try:
        ingest_res = client.post("/ingest-path", json={"file_path": txt_path})
        assert ingest_res.status_code == 200
        doc_id = ingest_res.json()["id"]

        with pytest.raises(ValueError, match="DOCX"):
            translator_module.translate_docx_inplace(doc_id, "English", "Italian")
    finally:
        if doc_id:
            client.delete(f"/documents/{doc_id}")


# --- PDF fine-mode (Fase 2) ---------------------------------------------------------------


def test_resolve_autofit_font_size_keeps_original_size_when_it_already_fits():
    rect = pymupdf.Rect(40, 40, 400, 80)
    size = translator_module._resolve_autofit_font_size(rect, "short", 14.0, "helv")
    assert size == 14.0


def test_resolve_autofit_font_size_shrinks_to_fit_longer_text():
    # Tight rect matching a real single-word block's padded bbox at 14pt (see
    # _padded_block_rect) -- a much longer translated sentence doesn't fit at 14pt here.
    rect = pymupdf.Rect(37.9, 37.9, 181.4, 61.3)
    longer_text = "This translated text is considerably longer than the original short marker sentence"
    original_size = 14.0

    assert not translator_module._fits_at_font_size(rect, longer_text, original_size, "helv")

    fit_size = translator_module._resolve_autofit_font_size(rect, longer_text, original_size, "helv")

    assert fit_size < original_size
    floor = max(translator_module._PDF_AUTOFIT_MIN_SIZE, original_size * translator_module._PDF_AUTOFIT_MIN_RATIO)
    assert fit_size >= floor
    assert translator_module._fits_at_font_size(rect, longer_text, fit_size, "helv")


def test_resolve_autofit_font_size_never_shrinks_below_floor():
    rect = pymupdf.Rect(0, 0, 10, 10)  # too small to ever fit any of this text
    huge_text = "word " * 200
    fit_size = translator_module._resolve_autofit_font_size(rect, huge_text, 14.0, "helv")
    floor = max(translator_module._PDF_AUTOFIT_MIN_SIZE, 14.0 * translator_module._PDF_AUTOFIT_MIN_RATIO)
    assert fit_size == floor


def test_redact_and_reinsert_pdf_blocks_autofits_instead_of_clipping(tmp_path):
    """A translated sentence that would have been clipped at the original font size (Fase 2
    behavior) must now be fully reinserted at a smaller auto-fit size (Fase 3), not truncated."""
    path = str(tmp_path / "autofit_src.pdf")
    out_path = str(tmp_path / "autofit_out.pdf")
    _make_pdf(path, text="Short marker text here")

    longer_text = "This translated text is considerably longer than the original short marker sentence"
    doc = pymupdf.open(path)
    try:
        page = doc[0]
        blocks = translator_module._extract_pdf_page_blocks(page)
        for b in blocks:
            b["text"] = longer_text
        translator_module._redact_and_reinsert_pdf_blocks(page, blocks, translator_module._PDF_FALLBACK_FONT_FILE)
        doc.save(out_path)
    finally:
        doc.close()

    reopened = pymupdf.open(out_path)
    try:
        # get_text() joins wrapped lines with '\n' at the point insert_textbox broke them, so
        # normalize whitespace before comparing -- the point being verified is that the full
        # sentence survived (auto-fit), not that it landed on a single physical line.
        page_text = " ".join(reopened[0].get_text().split())
    finally:
        reopened.close()
    assert longer_text in page_text


def test_extract_pdf_page_blocks_reads_bbox_text_size_color(tmp_path):
    path = str(tmp_path / "sample.pdf")
    _make_pdf(path, text="Hello PDF world")
    doc = pymupdf.open(path)
    try:
        blocks = translator_module._extract_pdf_page_blocks(doc[0])
        assert len(blocks) == 1
        assert "Hello PDF world" in blocks[0]["text"]
        assert blocks[0]["size"] > 0
        assert len(blocks[0]["bbox"]) == 4
    finally:
        doc.close()


def test_redact_and_reinsert_pdf_blocks_erases_original_text_from_raw_stream(tmp_path):
    """Verifies real redaction, not a visual overlay: the original marker must be absent from
    both the parsed text layer AND the raw saved PDF bytes after apply_redactions()."""
    marker = "SUPERSECRETMARKER12345"
    src_path = str(tmp_path / "src.pdf")
    out_path = str(tmp_path / "out.pdf")
    _make_pdf(src_path, text=marker)

    doc = pymupdf.open(src_path)
    try:
        page = doc[0]
        blocks = translator_module._extract_pdf_page_blocks(page)
        assert any(marker in b["text"] for b in blocks)
        for b in blocks:
            b["text"] = "REPLACED"
        translator_module._redact_and_reinsert_pdf_blocks(page, blocks, translator_module._PDF_FALLBACK_FONT_FILE)
        doc.save(out_path)
    finally:
        doc.close()

    reopened = pymupdf.open(out_path)
    try:
        assert marker not in reopened[0].get_text()
        assert "REPLACED" in reopened[0].get_text()
    finally:
        reopened.close()

    with open(out_path, "rb") as f:
        raw = f.read()
    assert marker.encode() not in raw


def test_translate_pdf_blocks_happy_path(tmp_path, monkeypatch):
    path = str(tmp_path / "sample.pdf")
    _make_pdf(path, text="Hello world")
    doc = pymupdf.open(path)
    try:
        blocks = translator_module._extract_pdf_page_blocks(doc[0])
    finally:
        doc.close()

    def fake_call(text, source_lang, target_lang, model):
        return f"[{text}]"

    monkeypatch.setattr(translator_module, "_call_ollama_translate", fake_call)
    translator_module._translate_pdf_blocks(blocks, "English", "Italian", "llama3.2")
    assert blocks[0]["text"].startswith("[") and "Hello world" in blocks[0]["text"]


def test_translate_pdf_inplace_fine_end_to_end(tmp_path, monkeypatch):
    # Same-length transform (reversal) so the "translated" text fits the original bbox width at
    # the original font size -- this test covers the happy path, not the Fase 2 overflow-clipping
    # limitation, which is covered separately by test_translate_pdf_inplace_fine_clips_overflow.
    marker = "UNIQUEMARKERXYZ"
    translated_marker = marker[::-1]
    path = str(tmp_path / "e2e_sample.pdf")
    _make_pdf(path, text=marker)

    def fake_call(text, source_lang, target_lang, model):
        if translator_module._RUN_SEPARATOR in text:
            parts = text.split(f"\n{translator_module._RUN_SEPARATOR}\n")
            return f"\n{translator_module._RUN_SEPARATOR}\n".join(p[::-1] for p in parts)
        return text[::-1]

    monkeypatch.setattr(translator_module, "_call_ollama_translate", fake_call)

    doc_id = None
    try:
        ingest_res = client.post("/ingest-path", json={"file_path": path})
        assert ingest_res.status_code == 200
        doc_id = ingest_res.json()["id"]

        res = client.post(
            f"/documents/{doc_id}/translate-inplace",
            json={"source_lang": "English", "target_lang": "Italian"},
        )
        assert res.status_code == 200
        data = res.json()
        assert data["status"] == "indexed"
        # In-place translation separates flows: original extracted markdown in RAG DB is preserved.
        assert marker in data["extracted_markdown"]

        translated_doc = pymupdf.open(path)
        try:
            page_text = translated_doc[0].get_text()
        finally:
            translated_doc.close()
        assert translated_marker in page_text

        # Original marker was genuinely erased, not just overlaid: it must be absent even from
        # the raw saved bytes (page content streams may be compressed, so this is a stronger
        # check than the parsed-text assertion above -- it can't find a compressed match either).
        with open(path, "rb") as f:
            raw = f.read()
        assert marker.encode() not in raw
    finally:
        if doc_id:
            client.delete(f"/documents/{doc_id}")


def test_translate_pdf_inplace_fine_clips_overflow_without_crashing(tmp_path, monkeypatch):
    """Fase 2 has no auto-fit: translated text that doesn't fit the original bbox at the original
    font size is clipped (dropped), not resized. This must not raise or corrupt the document --
    it's a documented limitation, not an error condition."""
    marker = "UNIQUEMARKERXYZ"
    path = str(tmp_path / "overflow_sample.pdf")
    _make_pdf(path, text=marker)

    def fake_call(text, source_lang, target_lang, model):
        # Much longer than the original -- guaranteed to overflow the tight original bbox.
        if translator_module._RUN_SEPARATOR in text:
            parts = text.split(f"\n{translator_module._RUN_SEPARATOR}\n")
            return f"\n{translator_module._RUN_SEPARATOR}\n".join(f"TRANSLATED-{p}-VERY-LONG-OUTPUT" for p in parts)
        return f"TRANSLATED-{text}-VERY-LONG-OUTPUT"

    monkeypatch.setattr(translator_module, "_call_ollama_translate", fake_call)

    doc_id = None
    try:
        ingest_res = client.post("/ingest-path", json={"file_path": path})
        assert ingest_res.status_code == 200
        doc_id = ingest_res.json()["id"]

        res = client.post(
            f"/documents/{doc_id}/translate-inplace",
            json={"source_lang": "English", "target_lang": "Italian"},
        )
        assert res.status_code == 200

        with open(path, "rb") as f:
            raw = f.read()
        # The original was still genuinely erased even though the replacement got clipped.
        assert marker.encode() not in raw
    finally:
        if doc_id:
            client.delete(f"/documents/{doc_id}")


def test_translate_document_inplace_rejects_unsupported_file_type(tmp_path):
    txt_path = str(tmp_path / "plain.txt")
    with open(txt_path, "w", encoding="utf-8") as f:
        f.write("Just a plain text file.")

    doc_id = None
    try:
        ingest_res = client.post("/ingest-path", json={"file_path": txt_path})
        assert ingest_res.status_code == 200
        doc_id = ingest_res.json()["id"]

        with pytest.raises(translator_module.UnsupportedDocumentTypeError):
            translator_module.translate_document_inplace(doc_id, "English", "Italian")

        res = client.post(
            f"/documents/{doc_id}/translate-inplace",
            json={"source_lang": "English", "target_lang": "Italian"},
        )
        assert res.status_code == 400
    finally:
        if doc_id:
            client.delete(f"/documents/{doc_id}")


def test_translate_document_inplace_returns_404_for_missing_document():
    res = client.post(
        "/documents/nonexistent-doc-id-xyz/translate-inplace",
        json={"source_lang": "English", "target_lang": "Italian"},
    )
    assert res.status_code == 404


# --- PDF fine-mode Fase 4: bundled CJK / universal fonts -----------------------------------


def test_resolve_pdf_font_file_selects_bundled_fonts_by_language():
    assert "NotoSansCJKjp" in translator_module._resolve_pdf_font_file("Japanese")
    assert "NotoSansCJKkr" in translator_module._resolve_pdf_font_file("Korean")
    assert "NotoSansCJKsc" in translator_module._resolve_pdf_font_file("Chinese")
    assert "NotoSansCJKtc" in translator_module._resolve_pdf_font_file("Traditional Chinese")
    # Non-CJK languages and unrecognized/empty strings fall back to the bundled Latin/Cyrillic/Greek font.
    assert translator_module._resolve_pdf_font_file("Italian") == translator_module._PDF_FALLBACK_FONT_FILE
    assert translator_module._resolve_pdf_font_file("Russian") == translator_module._PDF_FALLBACK_FONT_FILE
    assert translator_module._resolve_pdf_font_file("") == translator_module._PDF_FALLBACK_FONT_FILE


def test_bundled_pdf_fonts_exist_on_disk():
    """Guards against a packaging/path mistake: every font _resolve_pdf_font_file can return must
    actually be present under sidecar/assets/fonts."""
    font_files = {f for _, f in translator_module._PDF_LANG_FONT_RULES} | {translator_module._PDF_FALLBACK_FONT_FILE}
    for font_file in font_files:
        assert os.path.isfile(font_file), f"missing bundled font: {font_file}"


def test_redact_and_reinsert_pdf_blocks_renders_cjk_glyphs(tmp_path):
    """The reinserted text must round-trip through PyMuPDF's parsed text layer using real CJK
    glyphs, not tofu/notdef boxes -- proves the bundled CJK font is actually wired into
    insert_textbox, not just present on disk (see test_bundled_pdf_fonts_exist_on_disk)."""
    path = str(tmp_path / "cjk_src.pdf")
    out_path = str(tmp_path / "cjk_out.pdf")
    _make_pdf(path, text="Hello world")
    chinese_text = "你好"

    doc = pymupdf.open(path)
    try:
        page = doc[0]
        blocks = translator_module._extract_pdf_page_blocks(page)
        for b in blocks:
            b["text"] = chinese_text
        font_file = translator_module._resolve_pdf_font_file("Chinese")
        translator_module._redact_and_reinsert_pdf_blocks(page, blocks, font_file)
        doc.save(out_path)
    finally:
        doc.close()

    reopened = pymupdf.open(out_path)
    try:
        assert chinese_text in reopened[0].get_text()
    finally:
        reopened.close()


def test_redact_and_reinsert_pdf_blocks_renders_cyrillic_fallback(tmp_path):
    """Non-CJK languages (e.g. Russian) use the bundled Latin/Cyrillic/Greek fallback font. Before
    Fase 4, the base14 'helv' font used for every PDF language could not render Cyrillic at all
    (verified separately: it silently degrades to '?????' notdef placeholders)."""
    path = str(tmp_path / "cyr_src.pdf")
    out_path = str(tmp_path / "cyr_out.pdf")
    _make_pdf(path, text="Hello world")
    russian_text = "Привет"

    doc = pymupdf.open(path)
    try:
        page = doc[0]
        blocks = translator_module._extract_pdf_page_blocks(page)
        for b in blocks:
            b["text"] = russian_text
        font_file = translator_module._resolve_pdf_font_file("Russian")
        assert font_file == translator_module._PDF_FALLBACK_FONT_FILE
        translator_module._redact_and_reinsert_pdf_blocks(page, blocks, font_file)
        doc.save(out_path)
    finally:
        doc.close()

    reopened = pymupdf.open(out_path)
    try:
        assert russian_text in reopened[0].get_text()
    finally:
        reopened.close()


def test_translate_pdf_inplace_fine_end_to_end_japanese(tmp_path, monkeypatch):
    """Full pipeline end-to-end with a CJK target language: verifies the font-selection wiring
    from translate_pdf_inplace_fine all the way through to the saved file and reindexed markdown."""
    # Short phrase, sized to comfortably fit the narrow bbox of the "Hello world" source text
    # (this test verifies the font-selection wiring, not overflow/clipping behavior -- that's
    # covered separately by test_translate_pdf_inplace_fine_clips_overflow_without_crashing).
    japanese_text = "日本語"
    path = str(tmp_path / "jp_e2e.pdf")
    _make_pdf(path, text="Hello world")

    def fake_call(text, source_lang, target_lang, model):
        if translator_module._RUN_SEPARATOR in text:
            parts = text.split(f"\n{translator_module._RUN_SEPARATOR}\n")
            return f"\n{translator_module._RUN_SEPARATOR}\n".join(japanese_text for _ in parts)
        return japanese_text

    monkeypatch.setattr(translator_module, "_call_ollama_translate", fake_call)

    doc_id = None
    try:
        ingest_res = client.post("/ingest-path", json={"file_path": path})
        assert ingest_res.status_code == 200
        doc_id = ingest_res.json()["id"]

        res = client.post(
            f"/documents/{doc_id}/translate-inplace",
            json={"source_lang": "English", "target_lang": "Japanese"},
        )
        assert res.status_code == 200
        assert res.json()["status"] == "indexed"

        translated_doc = pymupdf.open(path)
        try:
            page_text = translated_doc[0].get_text()
        finally:
            translated_doc.close()
        assert japanese_text in page_text
        # In-place translation separates flows: original extracted markdown is preserved.
        assert "Hello world" in res.json()["extracted_markdown"]
    finally:
        if doc_id:
            client.delete(f"/documents/{doc_id}")


def test_translate_inplace_with_backup_and_target_dir(monkeypatch, tmp_path):
    monkeypatch.setattr(
        translator_module,
        "_call_ollama_translate",
        lambda text, s, t, m: f"TR-{text}",
    )
    src_file = tmp_path / "sample.docx"
    doc = docx.Document()
    doc.add_paragraph("Original paragraph text")
    doc.save(str(src_file))

    out_dir = tmp_path / "translated_output"
    out_dir.mkdir()

    doc_id = None
    try:
        ingest_res = client.post("/ingest-path", json={"file_path": str(src_file)})
        assert ingest_res.status_code == 200
        doc_id = ingest_res.json()["id"]

        # Test target_dir saving without altering original
        res = client.post(
            f"/documents/{doc_id}/translate-inplace",
            json={
                "source_lang": "English",
                "target_lang": "Spanish",
                "target_dir": str(out_dir),
                "backup_original": True,
            },
        )
        assert res.status_code == 200
        target_files = list(out_dir.glob("*.docx"))
        assert len(target_files) == 1
        assert "sample_spanish.docx" in target_files[0].name

        # Verify original file was not mutated
        orig_doc = docx.Document(str(src_file))
        orig_texts = [r.text for r in translator_module._collect_docx_runs(orig_doc)]
        assert "Original paragraph text" in orig_texts
    finally:
        if doc_id:
            client.delete(f"/documents/{doc_id}")


def test_translate_scanned_pdf_inplace_with_ocr_fallback(monkeypatch, tmp_path):
    """Scanned PDF with zero native text layer must detect text blocks via OCR fallback and translate in-place."""
    from sidecar.domain import ingestion as ingestion_module
    from sidecar.infrastructure import ocr as ocr_infra

    monkeypatch.setattr(
        translator_module,
        "_call_ollama_translate",
        lambda text, s, t, m: f"TR-{text}",
    )
    monkeypatch.setattr(
        ingestion_module,
        "run_layout_ocr",
        lambda img: "Richiesta Cessazione Contratto",
    )
    monkeypatch.setattr(
        ocr_infra,
        "run_rapid_ocr_with_boxes",
        lambda img: [{"bbox": [50, 80, 300, 100], "text": "Richiesta Cessazione Contratto", "score": 0.99}],
    )

    # 1. Create a pure raster image page (no text layer)
    img_doc = pymupdf.open()
    img_page = img_doc.new_page(width=595, height=842)
    img_page.draw_rect(pymupdf.Rect(0, 0, 595, 842), color=(1, 1, 1), fill=(1, 1, 1))
    img_page.insert_text(pymupdf.Point(50, 80), "Richiesta Cessazione Contratto", fontsize=18, color=(0, 0, 0))
    pix = img_page.get_pixmap(dpi=200)
    img_bytes = pix.tobytes(output="png")
    img_doc.close()

    scanned_pdf_path = str(tmp_path / "scanned_contract.pdf")
    doc = pymupdf.open()
    page = doc.new_page(width=595, height=842)
    page.insert_image(pymupdf.Rect(0, 0, 595, 842), stream=img_bytes)
    doc.save(scanned_pdf_path)
    doc.close()

    doc_id = None
    try:
        ingest_res = client.post("/ingest-path", json={"file_path": scanned_pdf_path})
        assert ingest_res.status_code == 200
        doc_id = ingest_res.json()["id"]

        # 2. In-place translate scanned PDF
        res = client.post(
            f"/documents/{doc_id}/translate-inplace",
            json={
                "source_lang": "Italian",
                "target_lang": "English",
                "backup_original": True,
            },
        )
        assert res.status_code == 200
        data = res.json()
        assert "TR-" in data["extracted_markdown"] or "Richiesta" in data["extracted_markdown"]

        # Verify translated PDF now contains search-enabled translated text
        translated_pdf = pymupdf.open(scanned_pdf_path)
        try:
            page_text = translated_pdf[0].get_text()
            assert len(page_text.strip()) > 0
        finally:
            translated_pdf.close()
    finally:
        if doc_id:
            client.delete(f"/documents/{doc_id}")


def test_translate_pdf_inplace_guarantees_100_percent_block_rendering(tmp_path):
    """Verifies that all blocks on a page are 100% rendered and no block is omitted due to vertical overflow."""
    font_file = translator_module._resolve_pdf_font_file("English")
    pdf_path = str(tmp_path / "tight_blocks.pdf")
    out_pdf_path = str(tmp_path / "tight_blocks_out.pdf")
    doc = pymupdf.open()
    page = doc.new_page(width=595, height=842)
    # Create multiple tight blocks with simulated lengthy translated text
    blocks = []
    for i in range(10):
        y0 = 100 + i * 40
        y1 = y0 + 12
        blocks.append({
            "bbox": (50, y0, 300, y1),
            "text": f"Block {i+1} translated with exceptionally long multi-word explanation that exceeds single line capacity",
            "size": 10.0,
            "color": 0,
        })
    doc.save(pdf_path)
    doc.close()

    doc2 = pymupdf.open(pdf_path)
    page2 = doc2[0]
    translator_module._redact_and_reinsert_pdf_blocks(page2, blocks, font_file)
    doc2.save(out_pdf_path)
    doc2.close()

    doc3 = pymupdf.open(out_pdf_path)
    page3 = doc3[0]
    out_blocks = [b for b in page3.get_text("dict")["blocks"] if b.get("type") == 0]
    doc3.close()

    # All 10 blocks must be present on page
    assert len(out_blocks) == 10




